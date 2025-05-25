"""
Generador de documentos Word - Versión Optimizada con manejo robusto de errores
"""

import os
import re
import threading
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_SECTION, WD_ORIENTATION
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from modules.watermark import WatermarkManager
except ImportError:
    WatermarkManager = None

from utils.logger import get_logger
from config.settings import DEFAULT_FORMAT

logger = get_logger("document_generator")

class DocumentGenerator:
    """Generador de documentos Word con manejo robusto de errores"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            logger.error("python-docx no está disponible. Instala con: pip install python-docx")
            
        self.formato_config = DEFAULT_FORMAT.copy()
        self.watermark_manager = WatermarkManager() if WatermarkManager else None
        
        # Configuración de generación
        self.max_image_size = (1920, 1080)
        self.supported_image_formats = ['.png', '.jpg', '.jpeg', '.gif', '.bmp']
        
        logger.info("DocumentGenerator inicializado")

    def generar_documento_async(self, app_instance):
        """Genera el documento profesional en un hilo separado con manejo robusto"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("❌ Error", 
                "python-docx no está disponible.\n"
                "Instala la dependencia con:\n"
                "pip install python-docx")
            return
        
        def generar():
            try:
                # Validar datos antes de generar
                if not self._validar_datos_proyecto(app_instance):
                    return
                
                # Inicializar progreso
                self._inicializar_progreso(app_instance)
                
                # Crear documento
                doc = Document()
                logger.info("Iniciando generación de documento")
                
                # Configurar documento base
                self.configurar_documento_completo(doc, app_instance)
                self._actualizar_progreso(app_instance, 0.1, "Configuración aplicada")
                
                # Generar secciones según configuración
                self._generar_secciones_documento(doc, app_instance)
                
                # Finalizar y guardar
                self._guardar_documento_final(doc, app_instance)
                
            except PermissionError as e:
                self._manejar_error(app_instance, "Error de permisos", 
                    "No se tiene permiso para escribir en la ubicación seleccionada")
            except FileNotFoundError as e:
                self._manejar_error(app_instance, "Archivo no encontrado", 
                    f"No se pudo encontrar un archivo necesario: {str(e)}")
            except Exception as e:
                self._manejar_error(app_instance, "Error inesperado", str(e))
        
        # Ejecutar en hilo separado
        thread = threading.Thread(target=generar, daemon=True)
        thread.start()
    
    def _validar_datos_proyecto(self, app_instance) -> bool:
        """Valida que los datos del proyecto son suficientes para generar"""
        try:
            # Verificar información básica
            if not hasattr(app_instance, 'proyecto_data'):
                messagebox.showerror("❌ Error", "No hay datos de proyecto disponibles")
                return False
            
            # Verificar al menos un campo de información general
            campos_basicos = ['titulo', 'estudiantes', 'institucion']
            tiene_info_basica = False
            
            for campo in campos_basicos:
                if (campo in app_instance.proyecto_data and 
                    hasattr(app_instance.proyecto_data[campo], 'get') and
                    app_instance.proyecto_data[campo].get().strip()):
                    tiene_info_basica = True
                    break
            
            if not tiene_info_basica:
                respuesta = messagebox.askyesno("⚠️ Información Incompleta",
                    "No hay información básica del proyecto.\n"
                    "¿Desea generar el documento de todas formas?")
                return respuesta
            
            # Verificar contenido
            if hasattr(app_instance, 'content_texts'):
                tiene_contenido = any(
                    text_widget.get("1.0", "end").strip() 
                    for text_widget in app_instance.content_texts.values()
                    if hasattr(text_widget, 'get')
                )
                
                if not tiene_contenido:
                    respuesta = messagebox.askyesno("⚠️ Sin Contenido",
                        "No hay contenido en las secciones.\n"
                        "¿Desea generar un documento en blanco?")
                    return respuesta
            
            return True
            
        except Exception as e:
            logger.error(f"Error validando datos del proyecto: {e}")
            return False
    
    def _generar_secciones_documento(self, doc, app_instance):
        """Genera todas las secciones del documento según configuración"""
        try:
            progress_step = 0.7 / 7  # 70% dividido entre las secciones principales
            current_progress = 0.2
            
            # Portada
            if getattr(app_instance, 'incluir_portada', ctk.BooleanVar()).get():
                self.crear_portada_profesional(doc, app_instance)
                current_progress += progress_step
                self._actualizar_progreso(app_instance, current_progress, "Portada creada")
            
            # Agradecimientos
            if getattr(app_instance, 'incluir_agradecimientos', ctk.BooleanVar()).get():
                self._crear_agradecimientos(doc, app_instance)
                current_progress += progress_step
                self._actualizar_progreso(app_instance, current_progress, "Agradecimientos agregados")
            
            # Resumen
            if self._seccion_tiene_contenido(app_instance, 'resumen'):
                self._crear_seccion_resumen(doc, app_instance)
                current_progress += progress_step
                self._actualizar_progreso(app_instance, current_progress, "Resumen agregado")
            
            # Índice
            if getattr(app_instance, 'incluir_indice', ctk.BooleanVar()).get():
                self.crear_indice_profesional(doc, app_instance)
                current_progress += progress_step
                self._actualizar_progreso(app_instance, current_progress, "Índice creado")
            
            # Contenido principal dinámico
            self.crear_contenido_dinamico_mejorado(doc, app_instance)
            current_progress += progress_step * 2
            self._actualizar_progreso(app_instance, current_progress, "Contenido principal generado")
            
            # Referencias
            if hasattr(app_instance, 'referencias') and app_instance.referencias:
                self.crear_referencias_profesionales(doc, app_instance)
                current_progress += progress_step
                self._actualizar_progreso(app_instance, current_progress, "Referencias agregadas")
            
        except Exception as e:
            logger.error(f"Error generando secciones: {e}")
            raise
    
    def _crear_agradecimientos(self, doc, app_instance):
        """Crea la sección de agradecimientos"""
        contenido_default = """Agradezco sinceramente a todas las personas que hicieron posible la realización de este proyecto académico.

A mis tutores y profesores por su guía y conocimientos compartidos.

A mi familia y amigos por su apoyo incondicional durante todo el proceso.

A la institución educativa por brindar los recursos y el ambiente necesario para el desarrollo de esta investigación."""
        
        contenido = contenido_default
        
        # Buscar contenido personalizado
        if (hasattr(app_instance, 'content_texts') and 
            'agradecimientos' in app_instance.content_texts):
            contenido_custom = app_instance.content_texts['agradecimientos'].get("1.0", "end").strip()
            if contenido_custom:
                contenido = contenido_custom
        
        self.crear_seccion_profesional(
            doc, "AGRADECIMIENTOS", contenido, app_instance, 
            nivel=1, aplicar_sangria_parrafos=False
        )
    
    def _crear_seccion_resumen(self, doc, app_instance):
        """Crea la sección de resumen"""
        if 'resumen' in app_instance.content_texts:
            contenido_resumen = app_instance.content_texts['resumen'].get("1.0", "end").strip()
            if contenido_resumen:
                contenido_normalizado = self.normalizar_parrafos(contenido_resumen)
                self.crear_seccion_profesional(
                    doc, "RESUMEN", contenido_normalizado, app_instance, 
                    nivel=1, aplicar_sangria_parrafos=False
                )
    
    def _seccion_tiene_contenido(self, app_instance, seccion_id) -> bool:
        """Verifica si una sección tiene contenido"""
        try:
            if (hasattr(app_instance, 'content_texts') and 
                seccion_id in app_instance.content_texts):
                contenido = app_instance.content_texts[seccion_id].get("1.0", "end").strip()
                return len(contenido) > 10
            return False
        except:
            return False
    
    def _inicializar_progreso(self, app_instance):
        """Inicializa la barra de progreso"""
        try:
            if hasattr(app_instance, 'progress'):
                app_instance.progress.set(0)
                if hasattr(app_instance.progress, 'start'):
                    app_instance.progress.start()
        except Exception as e:
            logger.warning(f"No se pudo inicializar progreso: {e}")
    
    def _actualizar_progreso(self, app_instance, valor, mensaje=""):
        """Actualiza la barra de progreso con mensaje"""
        try:
            if hasattr(app_instance, 'progress'):
                app_instance.progress.set(valor)
            
            if mensaje and hasattr(app_instance, 'validation_text'):
                app_instance.validation_text.delete("1.0", "end")
                app_instance.validation_text.insert("1.0", f"🔄 Generando documento...\n\n{mensaje}")
                
            logger.info(f"Progreso: {int(valor*100)}% - {mensaje}")
            
        except Exception as e:
            logger.warning(f"Error actualizando progreso: {e}")
    
    def _manejar_error(self, app_instance, titulo, mensaje):
        """Maneja errores durante la generación"""
        try:
            if hasattr(app_instance, 'progress'):
                app_instance.progress.stop()
                app_instance.progress.set(0)
            
            if hasattr(app_instance, 'validation_text'):
                app_instance.validation_text.delete("1.0", "end")
                app_instance.validation_text.insert("1.0", f"❌ Error en generación:\n\n{mensaje}")
            
            logger.error(f"{titulo}: {mensaje}")
            messagebox.showerror(f"❌ {titulo}", mensaje)
            
        except Exception as e:
            logger.error(f"Error manejando error: {e}")
    
    def _guardar_documento_final(self, doc, app_instance):
        """Guarda el documento final con validaciones"""
        try:
            # Sugerir nombre de archivo
            nombre_sugerido = self._generar_nombre_archivo(app_instance)
            
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[
                    ("Documentos Word", "*.docx"),
                    ("Todos los archivos", "*.*")
                ],
                title="Guardar Proyecto Académico",
                initialname=nombre_sugerido
            )
            
            if not filename:
                self._finalizar_progreso(app_instance, False)
                return
            
            # Validar extensión
            if not filename.lower().endswith('.docx'):
                filename += '.docx'
            
            # Verificar permisos de escritura
            try:
                test_path = Path(filename).parent / "test_write.tmp"
                test_path.touch()
                test_path.unlink()
            except Exception:
                raise PermissionError("No se tiene permiso de escritura en la ubicación seleccionada")
            
            # Guardar documento
            doc.save(filename)
            
            # Verificar que se guardó correctamente
            if not os.path.exists(filename) or os.path.getsize(filename) == 0:
                raise IOError("El documento no se guardó correctamente")
            
            self._finalizar_progreso(app_instance, True)
            self.mostrar_mensaje_exito(filename, app_instance)
            
            logger.info(f"Documento guardado exitosamente: {filename}")
            
        except Exception as e:
            logger.error(f"Error guardando documento: {e}")
            raise
    
    def _generar_nombre_archivo(self, app_instance) -> str:
        """Genera un nombre de archivo sugerido basado en el proyecto"""
        try:
            # Intentar usar el título del proyecto
            if (hasattr(app_instance, 'proyecto_data') and 
                'titulo' in app_instance.proyecto_data):
                titulo = app_instance.proyecto_data['titulo'].get().strip()
                if titulo:
                    # Limpiar caracteres no válidos para nombres de archivo
                    titulo_limpio = re.sub(r'[<>:"/\\|?*]', '', titulo)
                    titulo_limpio = titulo_limpio.replace('\n', ' ').strip()
                    if len(titulo_limpio) > 50:
                        titulo_limpio = titulo_limpio[:50] + "..."
                    return f"{titulo_limpio}.docx"
            
            # Nombre por defecto con timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            return f"proyecto_academico_{timestamp}.docx"
            
        except Exception as e:
            logger.warning(f"Error generando nombre de archivo: {e}")
            return f"proyecto_academico_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    def _finalizar_progreso(self, app_instance, exitoso=True):
        """Finaliza la barra de progreso"""
        try:
            if hasattr(app_instance, 'progress'):
                if hasattr(app_instance.progress, 'stop'):
                    app_instance.progress.stop()
                app_instance.progress.set(1 if exitoso else 0)
        except Exception as e:
            logger.warning(f"Error finalizando progreso: {e}")

    def normalizar_parrafos(self, contenido):
        """Normaliza párrafos con validación mejorada"""
        if not isinstance(contenido, str):
            return ""
        
        texto = contenido.strip()
        if not texto:
            return ""
        
        # Si ya hay dobles saltos, mantener estructura
        if '\n\n' in texto:
            return texto
        
        # Convertir saltos simples en dobles
        lineas = texto.split('\n')
        parrafos = [linea.strip() for linea in lineas if linea.strip()]
        
        return '\n\n'.join(parrafos)

    def generar_documento_async(self, app_instance):
        """Genera el documento profesional en un hilo separado"""
        def generar():
            try:
                app_instance.progress.set(0)
                app_instance.progress.start()
                
                doc = Document()
                
                self.configurar_documento_completo(doc, app_instance)
                app_instance.progress.set(0.1)
                
                # Portada
                if app_instance.incluir_portada.get():
                    self.crear_portada_profesional(doc, app_instance)
                    app_instance.progress.set(0.2)
                
                # Agradecimientos
                if app_instance.incluir_agradecimientos.get():
                    contenido_agradecimientos = "(Agregar agradecimientos personalizados aquí)"
                    contenido_agradecimientos = self.normalizar_parrafos(contenido_agradecimientos)
                    self.crear_seccion_profesional(
                        doc, "AGRADECIMIENTOS", contenido_agradecimientos, app_instance, nivel=1, aplicar_sangria_parrafos=False
                    )
                    app_instance.progress.set(0.3)
                
                # Resumen
                if 'resumen' in app_instance.secciones_activas and 'resumen' in app_instance.content_texts:
                    contenido_resumen = app_instance.content_texts['resumen'].get("1.0", "end")
                    contenido_resumen = self.normalizar_parrafos(contenido_resumen)
                    self.crear_seccion_profesional(
                        doc, "RESUMEN", contenido_resumen, app_instance, nivel=1, aplicar_sangria_parrafos=False
                    )
                    app_instance.progress.set(0.4)
                
                # Índice
                if app_instance.incluir_indice.get():
                    self.crear_indice_profesional(doc, app_instance)
                    app_instance.progress.set(0.5)
                
                # Contenido principal dinámico
                self.crear_contenido_dinamico_mejorado(doc, app_instance)
                app_instance.progress.set(0.8)
                
                # Referencias
                self.crear_referencias_profesionales(doc, app_instance)
                app_instance.progress.set(0.9)
                
                # Guardar documento
                filename = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word documents", "*.docx")],
                    title="Guardar Proyecto Académico Profesional"
                )
                
                if filename:
                    doc.save(filename)
                    app_instance.progress.stop()
                    app_instance.progress.set(1)
                    self.mostrar_mensaje_exito(filename, app_instance)
                else:
                    app_instance.progress.stop()
                    app_instance.progress.set(0)
            
            except Exception as e:
                app_instance.progress.stop()
                app_instance.progress.set(0)
                messagebox.showerror("❌ Error", f"Error al generar documento:\n{str(e)}")
        
        thread = threading.Thread(target=generar)
        thread.daemon = True
        thread.start()
    
    def configurar_documento_completo(self, doc, app_instance):
        """Configura el documento con validación de errores"""
        try:
            # Configurar márgenes
            for section in doc.sections:
                margen_cm = app_instance.formato_config.get('margen', 2.54)
                margen_inches = margen_cm / 2.54
                
                section.top_margin = Inches(margen_inches)
                section.bottom_margin = Inches(margen_inches)
                section.left_margin = Inches(margen_inches)
                section.right_margin = Inches(margen_inches)
                
                # Configurar encabezado
                self.configurar_encabezado_marca_agua(section, app_instance)
            
            # Configurar estilos
            self.configurar_estilos_profesionales(doc, app_instance)
            
        except Exception as e:
            logger.error(f"Error configurando documento: {e}")
            # Continuar con configuración básica
            self._configurar_documento_basico(doc)
    def _configurar_documento_basico(self, doc):
        """Configuración básica de respaldo"""
        try:
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Estilo normal básico
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
        except Exception as e:
            logger.error(f"Error en configuración básica: {e}")
    def configurar_encabezado_marca_agua(self, section, app_instance):
        """Configura encabezado con manejo robusto de errores"""
        try:
            section.different_first_page_header_footer = True
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.25)
            
            ruta_encabezado = self.obtener_ruta_imagen("encabezado", app_instance)
            
            if ruta_encabezado and os.path.exists(ruta_encabezado):
                self._configurar_encabezado_con_imagen(section, ruta_encabezado, app_instance)
            else:
                self._configurar_encabezado_simple(section, app_instance)
                
        except Exception as e:
            logger.warning(f"Error configurando encabezado: {e}")
            self._configurar_encabezado_simple(section, app_instance)
    def _configurar_encabezado_con_imagen(self, section, ruta_imagen, app_instance):
        """Configura encabezado con imagen"""
        try:
            header = section.header
            
            # Limpiar encabezado existente
            for para in header.paragraphs:
                p = para._element
                p.getparent().remove(p)
            
            # Agregar imagen
            header_para = header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_para.add_run()
            
            # Validar formato de imagen
            if not any(ruta_imagen.lower().endswith(ext) for ext in self.supported_image_formats):
                logger.warning(f"Formato de imagen no soportado: {ruta_imagen}")
                raise ValueError("Formato de imagen no soportado")
            
            # Agregar imagen con tamaño controlado
            run.add_picture(ruta_imagen, width=Cm(20))
            
            if self.watermark_manager:
                try:
                    header_pic = run.element.xpath('.//pic:pic')[0] if run.element.xpath('.//pic:pic') else None
                    if header_pic:
                        self.watermark_manager.configurar_imagen_detras_texto(
                            header_pic, 
                            self.watermark_manager.header_config
                        )
                except Exception as e:
                    logger.warning(f"Error aplicando marca de agua: {e}")
            
        except Exception as e:
            logger.warning(f"Error configurando encabezado con imagen: {e}")
            raise
    
    def _configurar_encabezado_simple(self, section, app_instance):
        """Configura encabezado simple con texto"""
        try:
            header = section.header
            
            # Limpiar encabezado existente
            for para in header.paragraphs:
                p = para._element
                p.getparent().remove(p)
            
            # Agregar texto
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Obtener nombre de institución
            institucion = "INSTITUCIÓN EDUCATIVA"
            if (hasattr(app_instance, 'proyecto_data') and 
                'institucion' in app_instance.proyecto_data):
                institucion_custom = app_instance.proyecto_data['institucion'].get()
                if institucion_custom:
                    institucion = institucion_custom.upper()
            
            run = p.add_run(institucion)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            
        except Exception as e:
            logger.error(f"Error configurando encabezado simple: {e}")
    def _configurar_encabezado_simple(self, section, app_instance):
        try:
            header = section.header
            for para in header.paragraphs:
                p = para._element
                p.getparent().remove(p)
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            institucion = app_instance.proyecto_data.get('institucion', None)
            if institucion and hasattr(institucion, 'get'):
                texto = institucion.get() or "INSTITUCIÓN EDUCATIVA"
            else:
                texto = "INSTITUCIÓN EDUCATIVA"
            run = p.add_run(texto.upper())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
        except Exception:
            pass
    
    def configurar_estilos_profesionales(self, doc, app_instance):
        style = doc.styles['Normal']
        style.font.name = app_instance.formato_config['fuente_texto']
        style.font.size = Pt(app_instance.formato_config['tamaño_texto'])
        style.font.color.rgb = RGBColor(0, 0, 0)
        if app_instance.formato_config['interlineado'] == 1.0:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        elif app_instance.formato_config['interlineado'] == 1.5:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        if app_instance.formato_config['justificado']:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.space_after = Pt(0)
        try:
            body_style = doc.styles.add_style('BodyTextIndent', WD_STYLE_TYPE.PARAGRAPH)
            body_style.base_style = doc.styles['Normal']
            body_style.font.name = app_instance.formato_config['fuente_texto']
            body_style.font.size = Pt(app_instance.formato_config['tamaño_texto'])
            if app_instance.formato_config['sangria']:
                body_style.paragraph_format.first_line_indent = Inches(0)
        except:
            if 'BodyTextIndent' in doc.styles:
                body_style = doc.styles['BodyTextIndent']
                body_style.font.name = app_instance.formato_config['fuente_texto']
                body_style.font.size = Pt(app_instance.formato_config['tamaño_texto'])
                if app_instance.formato_config['sangria']:
                    body_style.paragraph_format.first_line_indent = Inches(0.5)
        for i in range(1, 7):
            heading_name = f'Heading {i}'
            if heading_name in doc.styles:
                heading_style = doc.styles[heading_name]
            else:
                try:
                    heading_style = doc.styles.add_style(heading_name, WD_STYLE_TYPE.PARAGRAPH)
                except:
                    continue
            heading_style.font.name = app_instance.formato_config['fuente_titulo']
            heading_style.font.size = Pt(app_instance.formato_config['tamaño_titulo'])
            heading_style.paragraph_format.page_break_before = True
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(12)
            heading_style.paragraph_format.keep_with_next = True
            heading_style.paragraph_format.outline_level = i - 1
            heading_style.paragraph_format.first_line_indent = Inches(0)
    
    def crear_portada_profesional(self, doc, app_instance):
        ruta_insignia = self.obtener_ruta_imagen("insignia", app_instance)
        if ruta_insignia and os.path.exists(ruta_insignia):
            try:
                if hasattr(self, 'watermark_manager'):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Inches(0)
                    run = p.add_run()
                    run.add_picture(ruta_insignia, height=self.watermark_manager.logo_config['height'])
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(ruta_insignia, width=Inches(1.5))
            except Exception:
                pass
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(app_instance.proyecto_data['institucion'].get().upper())
        run.bold = True
        run.font.name = app_instance.formato_config['fuente_titulo']
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(f'"{app_instance.proyecto_data["titulo"].get()}"')
        run.bold = True
        run.font.name = app_instance.formato_config['fuente_titulo']
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 0, 0)
        info_fields = [
            ('ciclo', 'Ciclo'),
            ('curso', 'Curso'), 
            ('enfasis', 'Énfasis'),
            ('area', 'Área de Desarrollo'),
            ('categoria', 'Categoría'),
            ('director', 'Director'),
            ('responsable', 'Responsable')
        ]
        for field, label in info_fields:
            if field in app_instance.proyecto_data and app_instance.proyecto_data[field].get().strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                label_run = p.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.name = app_instance.formato_config['fuente_texto']
                label_run.font.size = Pt(14)
                label_run.font.color.rgb = RGBColor(0, 0, 0)
                valor_original = app_instance.proyecto_data[field].get()
                if field == 'responsable' and ',' in valor_original:
                    responsables = [resp.strip() for resp in valor_original.split(',')]
                    if len(responsables) == 1:
                        valor_formateado = responsables[0]
                    elif len(responsables) == 2:
                        valor_formateado = f"{responsables[0]} y {responsables[1]}"
                    else:
                        todos_menos_ultimo = ", ".join(responsables[:-1])
                        valor_formateado = f"{todos_menos_ultimo} y {responsables[-1]}"
                    value_run = p.add_run(valor_formateado)
                else:
                    value_run = p.add_run(valor_original)
                value_run.font.name = app_instance.formato_config['fuente_texto']
                value_run.font.size = Pt(12)
                value_run.font.color.rgb = RGBColor(0, 0, 0)
        if app_instance.proyecto_data['estudiantes'].get():
            self._agregar_lista_personas(doc, "Estudiantes", 
                                    app_instance.proyecto_data['estudiantes'].get(), 
                                    app_instance, alineacion='izquierda')
        if app_instance.proyecto_data['tutores'].get():
            self._agregar_lista_personas(doc, "Tutores", 
                                    app_instance.proyecto_data['tutores'].get(), 
                                    app_instance, alineacion='izquierda')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        year_label = p.add_run("Año: ")
        year_label.bold = True
        year_label.font.name = app_instance.formato_config['fuente_texto']
        year_label.font.size = Pt(14)
        year_label.font.color.rgb = RGBColor(0, 0, 0)
        year_value = p.add_run(str(datetime.now().year))
        year_value.font.name = app_instance.formato_config['fuente_texto']
        year_value.font.size = Pt(12)
        year_value.font.color.rgb = RGBColor(0, 0, 0)
        doc.add_page_break()

    def _agregar_lista_personas(self, doc, titulo, personas_str, app_instance, alineacion='centro'):
        p = doc.add_paragraph()
        if alineacion == 'izquierda':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = p.add_run(f"{titulo}: ")
        title_run.bold = True
        title_run.font.name = app_instance.formato_config['fuente_texto']
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        personas = []
        for persona in personas_str.split(','):
            persona_limpia = persona.strip()
            if persona_limpia:
                personas.append(persona_limpia)
        if len(personas) == 0:
            personas_run = p.add_run("")
        elif len(personas) == 1:
            personas_run = p.add_run(personas[0])
        elif len(personas) == 2:
            personas_run = p.add_run(f"{personas[0]} y {personas[1]}")
        else:
            todos_menos_ultimo = ", ".join(personas[:-1])
            personas_run = p.add_run(f"{todos_menos_ultimo} y {personas[-1]}")
        personas_run.font.name = app_instance.formato_config['fuente_texto']
        personas_run.font.size = Pt(12)
        personas_run.font.color.rgb = RGBColor(0, 0, 0)
    
    def crear_indice_profesional(self, doc, app_instance):
        p = doc.add_heading('ÍNDICE', level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        doc.add_paragraph()
        instrucciones = """INSTRUCCIONES PARA GENERAR ÍNDICE AUTOMÁTICO:

    1. En Word, ir a la pestaña "Referencias"
    2. Hacer clic en "Tabla de contenido"  
    3. Seleccionar el estilo deseado
    4. El índice se generará automáticamente

    NOTA: Todos los títulos están configurados con niveles de esquema para facilitar la generación automática."""
        for linea in instrucciones.split('\n'):
            p = doc.add_paragraph(linea)
            p.paragraph_format.first_line_indent = Inches(0)
        doc.add_paragraph()
        p = doc.add_heading('TABLA DE ILUSTRACIONES', level=2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p = doc.add_paragraph("(Agregar manualmente si hay figuras, tablas o gráficos)")
        p.paragraph_format.first_line_indent = Inches(0)
        doc.add_page_break()
    
    def crear_contenido_dinamico_mejorado(self, doc, app_instance):
        capitulo_num = 0
        secciones_sin_sangria = [
            "RESUMEN", "PALABRAS CLAVE", "AGRADECIMIENTOS",
            "ÍNDICE", "TABLA DE ILUSTRACIONES", "REFERENCIAS"
        ]
        for seccion_id in app_instance.secciones_activas:
            if seccion_id in app_instance.secciones_disponibles:
                seccion = app_instance.secciones_disponibles[seccion_id]
                if seccion['capitulo']:
                    capitulo_num += 1
                    titulo = seccion['titulo']
                    titulo_limpio = re.sub(r'[^\w\s-]', '', titulo).strip()
                    p = doc.add_heading(titulo_limpio, level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Inches(0)
                    if capitulo_num > 1:
                        doc.add_page_break()
                    else:
                        doc.add_paragraph()
                else:
                    if seccion_id in app_instance.content_texts:
                        raw_content = app_instance.content_texts[seccion_id].get("1.0", "end")
                        contenido = self.normalizar_parrafos(raw_content)
                        if contenido:
                            titulo = seccion['titulo']
                            titulo_limpio = re.sub(r'[^\w\s-]', '', titulo).strip()
                            aplicar_sangria = not (titulo_limpio.upper() in secciones_sin_sangria)
                            self.crear_seccion_profesional(
                                doc, titulo_limpio.upper(), contenido, app_instance, nivel=2,
                                aplicar_sangria_parrafos=aplicar_sangria
                            )

    def crear_seccion_profesional(self, doc, titulo, contenido, app_instance, nivel=1, aplicar_sangria_parrafos=True):
        p = doc.add_heading(titulo, level=nivel)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        contenido_procesado = self.procesar_citas_mejorado(contenido.strip(), app_instance)
        parrafos = contenido_procesado.split('\n\n')
        for i, parrafo in enumerate(parrafos):
            texto = parrafo.strip()
            if texto:
                p = doc.add_paragraph(texto)
                aplicar_sangria = aplicar_sangria_parrafos
                if len(texto.split()) > 40:
                    aplicar_sangria = False
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(0)
                elif texto.startswith('\t') or texto.startswith('     '):
                    aplicar_sangria = False
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(0)
                elif any(texto.startswith(marca) for marca in [
                    '•', '-', '*', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.'
                ]):
                    aplicar_sangria = False
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(0)
                if aplicar_sangria and app_instance.formato_config.get('sangria', True):
                    p.paragraph_format.first_line_indent = Inches(0.5)
                else:
                    p.paragraph_format.first_line_indent = Inches(0)
                p.style = doc.styles['Normal']
        doc.add_paragraph()
    
    def procesar_citas_mejorado(self, texto, app_instance):
        if hasattr(app_instance, 'citation_processor'):
            return app_instance.citation_processor.procesar_citas_avanzado(texto)
        def reemplazar_cita(match):
            cita_completa = match.group(0)
            contenido = cita_completa[6:-1]
            partes = contenido.split(':')
            if len(partes) >= 3:
                tipo, autor, año = partes[0], partes[1], partes[2]
                pagina = partes[3] if len(partes) > 3 else None
                if tipo == 'textual':
                    if pagina:
                        return f" ({autor}, {año}, p. {pagina})"
                    else:
                        return f" ({autor}, {año})"
                elif tipo == 'parafraseo':
                    return f" ({autor}, {año})"
                elif tipo == 'larga':
                    if pagina:
                        return f"\n\n     ({autor}, {año}, p. {pagina})\n\n"
                    else:
                        return f"\n\n     ({autor}, {año})\n\n"
                elif tipo == 'web':
                    return f" ({autor}, {año})"
                elif tipo == 'multiple':
                    return f" ({autor}, {año})"
                else:
                    return f" ({autor}, {año})"
            return cita_completa
        texto_procesado = re.sub(r'\[CITA:[^\]]+\]', reemplazar_cita, texto)
        return texto_procesado
    
    def crear_referencias_profesionales(self, doc, app_instance):
        if not app_instance.referencias:
            return
        p = doc.add_heading('REFERENCIAS', level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        referencias_ordenadas = sorted(app_instance.referencias, 
                                    key=lambda x: x['autor'].split(',')[0].strip())
        for ref in referencias_ordenadas:
            ref_text = self._formatear_referencia_apa(ref)
            p = doc.add_paragraph(ref_text)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            p.style = doc.styles['Normal']

    def _formatear_referencia_apa(self, ref):
        tipo = ref.get('tipo', 'Libro')
        autor = ref.get('autor', '')
        año = ref.get('año', '')
        titulo = ref.get('titulo', '')
        fuente = ref.get('fuente', '')
        if tipo == 'Libro':
            return f"{autor} ({año}). {titulo}. {fuente}."
        elif tipo == 'Artículo':
            return f"{autor} ({año}). {titulo}. {fuente}."
        elif tipo == 'Web':
            return f"{autor} ({año}). {titulo}. Recuperado de {fuente}"
        elif tipo == 'Tesis':
            return f"{autor} ({año}). {titulo} [Tesis]. {fuente}."
        else:
            return f"{autor} ({año}). {titulo}. {fuente}."
    
    def obtener_ruta_imagen(self, tipo, app_instance):
        if tipo == "encabezado":
            return (getattr(app_instance, 'encabezado_personalizado', None) or 
                   getattr(app_instance, 'ruta_encabezado', None))
        elif tipo == "insignia":
            return (getattr(app_instance, 'insignia_personalizada', None) or 
                   getattr(app_instance, 'ruta_insignia', None))
        return None

    def mostrar_mensaje_exito(self, filename, app_instance):
        """Muestra mensaje de éxito mejorado"""
        try:
            # Calcular estadísticas del documento
            file_size = os.path.getsize(filename)
            file_size_mb = file_size / (1024 * 1024)
            
            # Actualizar área de validación
            if hasattr(app_instance, 'validation_text'):
                mensaje_completo = (
                    f"🎉 ¡DOCUMENTO GENERADO EXITOSAMENTE!\n\n"
                    f"📄 Archivo: {os.path.basename(filename)}\n"
                    f"📍 Ubicación: {filename}\n"
                    f"📊 Tamaño: {file_size_mb:.2f} MB\n"
                    f"⏰ Generado: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n\n"
                    f"✅ CARACTERÍSTICAS APLICADAS:\n"
                    f"   • Formato académico profesional\n"
                    f"   • Encabezados personalizados\n"
                    f"   • Niveles de esquema para índice automático\n"
                    f"   • Formato de citas APA\n"
                    f"   • Referencias bibliográficas ordenadas\n"
                    f"   • Márgenes y tipografía estándar\n\n"
                    f"📋 PRÓXIMOS PASOS:\n"
                    f"   1. Abrir el documento en Microsoft Word\n"
                    f"   2. Ir a Referencias > Tabla de contenido\n"
                    f"   3. Seleccionar estilo automático\n"
                    f"   4. Revisar y ajustar según necesidades\n\n"
                    f"🚀 ¡Tu proyecto académico está listo!"
                )
                
                app_instance.validation_text.delete("1.0", "end")
                app_instance.validation_text.insert("1.0", mensaje_completo)
            
            # Mostrar diálogo de éxito
            messagebox.showinfo("🎉 ¡Generación Exitosa!", 
                f"Documento generado correctamente:\n\n"
                f"📄 {os.path.basename(filename)}\n"
                f"📊 Tamaño: {file_size_mb:.2f} MB\n\n"
                f"El documento incluye:\n"
                f"✓ Formato académico profesional\n"
                f"✓ Estructura lista para índice automático\n"
                f"✓ Referencias en formato APA\n"
                f"✓ Configuración de página estándar")
            
        except Exception as e:
            logger.warning(f"Error mostrando mensaje de éxito: {e}")
            messagebox.showinfo("✅ Éxito", f"Documento generado: {filename}")

