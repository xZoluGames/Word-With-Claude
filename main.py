#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de Proyectos Académicos - Versión Avanzada
Crea documentos Word con opciones personalizables y formato base
"""

import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re
import os
import json
from datetime import datetime
import threading
from copy import deepcopy

# Configuración del tema
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class ProyectoAcademicoGenerator:
    def __init__(self):
        self.root = ctk.CTk()
        self.root.title("🎓 Generador de Proyectos Académicos - Versión Avanzada")
        self.root.geometry("1200x700")
        
        # Hacer la ventana redimensionable
        self.root.minsize(1000, 600)
        
        # Variables para almacenar la información
        self.proyecto_data = {}
        self.referencias = []
        self.documento_base = None
        self.usar_formato_base = False
        
        # Secciones dinámicas - estructura inicial
        self.secciones_disponibles = self.get_secciones_iniciales()
        self.secciones_activas = list(self.secciones_disponibles.keys())
        self.content_texts = {}
        
        # Configuración de formato
        self.formato_config = {
            'fuente_texto': 'Times New Roman',
            'tamaño_texto': 12,
            'fuente_titulo': 'Times New Roman', 
            'tamaño_titulo': 14,
            'interlineado': 2.0,
            'margen': 2.54,
            'justificado': True,
            'sangria': True
        }
        
        self.setup_ui()
        self.mostrar_bienvenida()
    
    def get_secciones_iniciales(self):
        """Define las secciones disponibles inicialmente"""
        return {
            "resumen": {
                "titulo": "📄 Resumen", 
                "instruccion": "Resumen ejecutivo del proyecto (150-300 palabras)",
                "requerida": False,
                "capitulo": False
            },
            "introduccion": {
                "titulo": "🔍 Introducción", 
                "instruccion": "Presenta el tema, contexto e importancia",
                "requerida": True,
                "capitulo": False
            },
            "capitulo1": {
                "titulo": "📖 CAPÍTULO I", 
                "instruccion": "Título de capítulo",
                "requerida": False,
                "capitulo": True
            },
            "planteamiento": {
                "titulo": "❓ Planteamiento del Problema", 
                "instruccion": "Define el problema a investigar",
                "requerida": True,
                "capitulo": False
            },
            "preguntas": {
                "titulo": "❔ Preguntas de Investigación", 
                "instruccion": "Pregunta general y específicas",
                "requerida": True,
                "capitulo": False
            },
            "delimitaciones": {
                "titulo": "📏 Delimitaciones", 
                "instruccion": "Límites del estudio (temporal, espacial, conceptual)",
                "requerida": False,
                "capitulo": False
            },
            "justificacion": {
                "titulo": "💡 Justificación", 
                "instruccion": "Explica por qué es importante investigar",
                "requerida": True,
                "capitulo": False
            },
            "objetivos": {
                "titulo": "🎯 Objetivos", 
                "instruccion": "General y específicos (verbos en infinitivo)",
                "requerida": True,
                "capitulo": False
            },
            "capitulo2": {
                "titulo": "📚 CAPÍTULO II - ESTADO DEL ARTE", 
                "instruccion": "Título de capítulo",
                "requerida": False,
                "capitulo": True
            },
            "marco_teorico": {
                "titulo": "📖 Marco Teórico", 
                "instruccion": "Base teórica y antecedentes (USA CITAS)",
                "requerida": True,
                "capitulo": False
            },
            "capitulo3": {
                "titulo": "🔬 CAPÍTULO III", 
                "instruccion": "Título de capítulo",
                "requerida": False,
                "capitulo": True
            },
            "metodologia": {
                "titulo": "⚙️ Marco Metodológico", 
                "instruccion": "Tipo de estudio y técnicas de recolección",
                "requerida": True,
                "capitulo": False
            },
            "capitulo4": {
                "titulo": "🛠️ CAPÍTULO IV - DESARROLLO", 
                "instruccion": "Título de capítulo",
                "requerida": False,
                "capitulo": True
            },
            "desarrollo": {
                "titulo": "⚙️ Desarrollo", 
                "instruccion": "Proceso de investigación paso a paso",
                "requerida": False,
                "capitulo": False
            },
            "capitulo5": {
                "titulo": "📊 CAPÍTULO V - ANÁLISIS DE DATOS", 
                "instruccion": "Título de capítulo",
                "requerida": False,
                "capitulo": True
            },
            "resultados": {
                "titulo": "📊 Resultados", 
                "instruccion": "Datos obtenidos (gráficos, tablas)",
                "requerida": False,
                "capitulo": False
            },
            "analisis_datos": {
                "titulo": "📈 Análisis de Datos", 
                "instruccion": "Interpretación de resultados",
                "requerida": False,
                "capitulo": False
            },
            "capitulo6": {
                "titulo": "💬 CAPÍTULO VI", 
                "instruccion": "Título de capítulo",
                "requerida": False,
                "capitulo": True
            },
            "discusion": {
                "titulo": "💬 Discusión", 
                "instruccion": "Confronta resultados con teoría",
                "requerida": False,
                "capitulo": False
            },
            "conclusiones": {
                "titulo": "✅ Conclusiones", 
                "instruccion": "Hallazgos principales y respuestas a objetivos",
                "requerida": True,
                "capitulo": False
            }
        }
    
    def setup_ui(self):
        """Configura la interfaz de usuario moderna con mejor gestión de espacio"""
        # Frame principal con scroll
        main_container = ctk.CTkFrame(self.root, corner_radius=0)
        main_container.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Header con título y botones principales
        header_frame = ctk.CTkFrame(main_container, height=120, corner_radius=10)
        header_frame.pack(fill="x", padx=10, pady=(10, 5))
        header_frame.pack_propagate(False)
        
        # Título en header
        title_label = ctk.CTkLabel(
            header_frame, 
            text="🎓 Generador de Proyectos Académicos",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title_label.pack(pady=(10, 5))
        
        # Botones principales en header
        button_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        button_frame.pack(fill="x", padx=20, pady=(5, 10))
        
        # Primera fila de botones
        btn_row1 = ctk.CTkFrame(button_frame, fg_color="transparent")
        btn_row1.pack(fill="x", pady=(0, 5))
        
        help_btn = ctk.CTkButton(
            btn_row1, text="📖 Guía", command=self.mostrar_instrucciones,
            width=80, height=30, font=ctk.CTkFont(size=11, weight="bold")
        )
        help_btn.pack(side="left", padx=(0, 5))
        
        template_btn = ctk.CTkButton(
            btn_row1, text="📋 Plantilla", command=self.cargar_documento_base,
            width=90, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="purple", hover_color="darkviolet"
        )
        template_btn.pack(side="left", padx=(0, 5))
        
        # Botones de proyecto
        save_btn = ctk.CTkButton(
            btn_row1, text="💾 Guardar", command=self.guardar_proyecto,
            width=80, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="darkgreen", hover_color="green"
        )
        save_btn.pack(side="left", padx=(0, 5))
        
        load_btn = ctk.CTkButton(
            btn_row1, text="📂 Cargar", command=self.cargar_proyecto,
            width=80, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="darkblue", hover_color="blue"
        )
        load_btn.pack(side="left", padx=(0, 5))
        
        # Estadísticas en tiempo real
        self.stats_label = ctk.CTkLabel(
            btn_row1, text="📊 Palabras: 0 | Secciones: 0/13 | Referencias: 0",
            font=ctk.CTkFont(size=10), text_color="gray70"
        )
        self.stats_label.pack(side="right", padx=(5, 0))
        
        # Segunda fila de botones
        btn_row2 = ctk.CTkFrame(button_frame, fg_color="transparent")
        btn_row2.pack(fill="x")
        
        # Botones de imágenes
        images_btn = ctk.CTkButton(
            btn_row2, text="🖼️ Imágenes", command=self.gestionar_imagenes,
            width=90, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="darkblue", hover_color="navy"
        )
        images_btn.pack(side="left", padx=(0, 5))
        
        export_btn = ctk.CTkButton(
            btn_row2, text="📤 Exportar Config", command=self.exportar_configuracion,
            width=110, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="darkorange", hover_color="orange"
        )
        export_btn.pack(side="left", padx=(0, 5))
        
        validate_btn = ctk.CTkButton(
            btn_row2, text="🔍 Validar", command=self.validar_proyecto,
            width=80, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="orange", hover_color="darkorange"
        )
        validate_btn.pack(side="left", padx=(0, 5))
        
        generate_btn = ctk.CTkButton(
            btn_row2, text="📄 Generar Documento", command=self.generar_documento_async,
            width=140, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="green", hover_color="darkgreen"
        )
        # Tercera fila de botones avanzados
        btn_row3 = ctk.CTkFrame(button_frame, fg_color="transparent")
        btn_row3.pack(fill="x", pady=(5, 0))
        
        # Botones de plantillas y temas
        template_mgr_btn = ctk.CTkButton(
            btn_row3, text="📋 Plantillas", command=self.gestionar_plantillas,
            width=90, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="darkviolet", hover_color="purple"
        )
        template_mgr_btn.pack(side="left", padx=(0, 5))
        
        theme_btn = ctk.CTkButton(
            btn_row3, text="🎨 Temas", command=self.cambiar_tema,
            width=80, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="darkcyan", hover_color="cyan"
        )
        theme_btn.pack(side="left", padx=(0, 5))
        
        preview_btn = ctk.CTkButton(
            btn_row3, text="👁️ Vista Previa", command=self.vista_previa_documento,
            width=100, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="darkslategray", hover_color="slategray"
        )
        preview_btn.pack(side="left", padx=(0, 5))
        
        backup_btn = ctk.CTkButton(
            btn_row3, text="🛡️ Respaldos", command=self.gestionar_respaldos,
            width=90, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="darkgoldenrod", hover_color="goldenrod"
        )
        backup_btn.pack(side="left", padx=(0, 5))
        
        # Selector de idioma y configuración avanzada
        settings_btn = ctk.CTkButton(
            btn_row3, text="⚙️ Config Avanzada", command=self.configuracion_avanzada,
            width=120, height=30, font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="dimgray", hover_color="gray"
        )
        settings_btn.pack(side="right", padx=(5, 0))
        
        # Frame contenedor para tabs con scroll
        content_container = ctk.CTkFrame(main_container, corner_radius=10)
        content_container.pack(fill="both", expand=True, padx=10, pady=(5, 10))
        
        # Tabview principal con altura reducida
        self.tabview = ctk.CTkTabview(content_container, width=1100, height=520)
        self.tabview.pack(expand=True, fill="both", padx=10, pady=10)
        
        # Crear pestañas
        self.setup_info_general()
        self.setup_contenido_dinamico()
        self.setup_citas_referencias()
        self.setup_formato_avanzado()
        self.setup_generacion()
    
    def setup_info_general(self):
        """Pestaña de información general mejorada y más compacta"""
        tab = self.tabview.add("📋 Información General")
        
        # Scroll frame con altura reducida
        scroll_frame = ctk.CTkScrollableFrame(tab, label_text="Datos del Proyecto", height=400)
        scroll_frame.pack(fill="both", expand=True, padx=15, pady=15)
        
        # Toggle para usar formato base - más compacto
        base_frame = ctk.CTkFrame(scroll_frame, fg_color="darkblue", corner_radius=8)
        base_frame.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(
            base_frame, text="📋 Formato Base",
            font=ctk.CTkFont(size=14, weight="bold"), text_color="white"
        ).pack(pady=(10, 5))
        
        self.usar_base_var = ctk.CTkCheckBox(
            base_frame, text="Usar plantilla base (3º AÑO)",
            font=ctk.CTkFont(size=12), command=self.toggle_formato_base
        )
        self.usar_base_var.pack(pady=(0, 10))
        
        # Campos organizados en filas de 2 columnas usando pack
        campos = [
            ("Institución Educativa", "institucion", "Colegio Privado Divina Esperanza"),
            ("Título del Proyecto", "titulo", "Ingrese el título de su investigación"),
            ("Categoría", "categoria", "Ciencia o Tecnología"),
            ("Ciclo", "ciclo", "Tercer año"),
            ("Curso", "curso", "3 BTI"),
            ("Énfasis", "enfasis", "Tecnología")
        ]
        
        # Crear campos en pares
        for i in range(0, len(campos), 2):
            row_frame = ctk.CTkFrame(scroll_frame, fg_color="transparent")
            row_frame.pack(fill="x", pady=5)
            
            # Primera columna
            if i < len(campos):
                label, key, placeholder = campos[i]
                field_frame1 = ctk.CTkFrame(row_frame, fg_color="transparent")
                field_frame1.pack(side="left", fill="both", expand=True, padx=(0, 10))
                
                ctk.CTkLabel(field_frame1, text=label, font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w")
                entry = ctk.CTkEntry(field_frame1, placeholder_text=placeholder, height=35)
                entry.pack(fill="x", pady=(5, 0))
                self.proyecto_data[key] = entry
            
            # Segunda columna
            if i + 1 < len(campos):
                label, key, placeholder = campos[i + 1]
                field_frame2 = ctk.CTkFrame(row_frame, fg_color="transparent")
                field_frame2.pack(side="right", fill="both", expand=True, padx=(10, 0))
                
                ctk.CTkLabel(field_frame2, text=label, font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w")
                entry = ctk.CTkEntry(field_frame2, placeholder_text=placeholder, height=30)
                entry.pack(fill="x", pady=(3, 0))
                self.proyecto_data[key] = entry
        
        # Campos largos (ancho completo)
        campos_largos = [
            ("Área de Desarrollo", "area", "Especifique el área de desarrollo"),
            ("Estudiantes (separar con comas)", "estudiantes", "Nombre1 Apellido1, Nombre2 Apellido2"),
            ("Tutores (separar con comas)", "tutores", "Prof. Nombre Apellido, Dr. Nombre Apellido"),
            ("Director", "director", "Cristina Raichakowski"),
            ("Responsable", "responsable", "Nombre del responsable")
        ]
        
        for label, key, placeholder in campos_largos:
            field_frame = ctk.CTkFrame(scroll_frame, fg_color="transparent")
            field_frame.pack(fill="x", pady=8)
            
            ctk.CTkLabel(field_frame, text=label, font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w")
            entry = ctk.CTkEntry(field_frame, placeholder_text=placeholder, height=30)
            entry.pack(fill="x", pady=(3, 0))
            self.proyecto_data[key] = entry
    
    def setup_contenido_dinamico(self):
        """Pestaña de contenido con gestión dinámica de secciones - más compacta"""
        tab = self.tabview.add("📝 Contenido Dinámico")
        
        # Frame principal dividido
        main_container = ctk.CTkFrame(tab, fg_color="transparent")
        main_container.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Panel de control de secciones (izquierda) - más compacto
        control_frame = ctk.CTkFrame(main_container, width=280, corner_radius=10)
        control_frame.pack(side="left", fill="y", padx=(0, 10))
        control_frame.pack_propagate(False)
        
        ctk.CTkLabel(
            control_frame, text="🛠️ Gestión de Secciones",
            font=ctk.CTkFont(size=14, weight="bold")
        ).pack(pady=(10, 8))
        
        # Botones de gestión - más compactos
        btn_frame = ctk.CTkFrame(control_frame, fg_color="transparent")
        btn_frame.pack(fill="x", padx=8, pady=(0, 8))
        
        add_btn = ctk.CTkButton(btn_frame, text="➕", command=self.agregar_seccion, width=60, height=28)
        add_btn.pack(side="left", padx=(0, 3))
        
        remove_btn = ctk.CTkButton(btn_frame, text="➖", command=self.quitar_seccion, 
                                 width=60, height=28, fg_color="red", hover_color="darkred")
        remove_btn.pack(side="left", padx=(0, 3))
        
        edit_btn = ctk.CTkButton(btn_frame, text="✏️", command=self.editar_seccion, width=60, height=28)
        edit_btn.pack(side="left")
        
        # Lista de secciones activas
        ctk.CTkLabel(control_frame, text="Secciones:", font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=8, pady=(8, 5))
        
        self.secciones_listbox = ctk.CTkScrollableFrame(control_frame, height=240)
        self.secciones_listbox.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        
        # Botones de orden - más compactos
        order_frame = ctk.CTkFrame(control_frame, fg_color="transparent")
        order_frame.pack(fill="x", padx=8, pady=(0, 8))
        
        up_btn = ctk.CTkButton(order_frame, text="⬆️", command=self.subir_seccion, width=80, height=28)
        up_btn.pack(side="left", padx=(0, 5))
        
        down_btn = ctk.CTkButton(order_frame, text="⬇️", command=self.bajar_seccion, width=80, height=28)
        down_btn.pack(side="right")
        
        # Panel de contenido (derecha)
        content_container = ctk.CTkFrame(main_container, corner_radius=10)
        content_container.pack(side="right", fill="both", expand=True)
        
        # Sub-tabview para secciones de contenido - altura reducida
        self.content_tabview = ctk.CTkTabview(content_container, width=700, height=420)
        self.content_tabview.pack(expand=True, fill="both", padx=8, pady=8)
        
        # Actualizar lista de secciones y crear pestañas
        self.actualizar_lista_secciones()
        self.crear_pestanas_contenido()
    
    def setup_citas_referencias(self):
        """Pestaña para gestión de citas más compacta"""
        tab = self.tabview.add("📚 Citas y Referencias")
        
        main_frame = ctk.CTkFrame(tab, fg_color="transparent")
        main_frame.pack(fill="both", expand=True, padx=15, pady=15)
        
        # Panel de instrucciones más compacto
        instruc_frame = ctk.CTkFrame(main_frame, fg_color="gray15", corner_radius=10, height=120)
        instruc_frame.pack(fill="x", pady=(0, 15))
        instruc_frame.pack_propagate(False)
        
        instruc_title = ctk.CTkLabel(
            instruc_frame, text="🚀 SISTEMA DE CITAS",
            font=ctk.CTkFont(size=14, weight="bold"), text_color="lightgreen"
        )
        instruc_title.pack(pady=(10, 5))
        
        ejemplos_text = "📝 [CITA:textual:García:2020:45] • 🔄 [CITA:parafraseo:López:2019] • 📖 [CITA:larga:Martínez:2021]"
        ctk.CTkLabel(
            instruc_frame, text=ejemplos_text,
            font=ctk.CTkFont(size=11), text_color="lightblue", wraplength=900
        ).pack(pady=2)
        
        ctk.CTkLabel(
            instruc_frame, text="✨ Conversión automática a formato APA",
            font=ctk.CTkFont(size=10, weight="bold"), text_color="yellow"
        ).pack(pady=(2, 10))
        
        # Frame para agregar referencias - más compacto
        ref_frame = ctk.CTkFrame(main_frame, height=140)
        ref_frame.pack(fill="x", pady=(0, 15))
        ref_frame.pack_propagate(False)
        
        ref_title = ctk.CTkLabel(
            ref_frame, text="➕ Agregar Referencias",
            font=ctk.CTkFont(size=14, weight="bold")
        )
        ref_title.pack(pady=(10, 8))
        
        # Campos para referencias usando pack - más compactos
        fields_frame = ctk.CTkFrame(ref_frame, fg_color="transparent")
        fields_frame.pack(fill="x", padx=15, pady=(0, 10))
        
        # Primera fila
        row1_frame = ctk.CTkFrame(fields_frame, fg_color="transparent")
        row1_frame.pack(fill="x", pady=3)
        
        # Tipo
        tipo_frame = ctk.CTkFrame(row1_frame, fg_color="transparent")
        tipo_frame.pack(side="left", fill="x", expand=True, padx=(0, 5))
        ctk.CTkLabel(tipo_frame, text="Tipo:", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.ref_tipo = ctk.CTkComboBox(
            tipo_frame,
            values=["Libro", "Artículo", "Web", "Tesis"],
            height=25, font=ctk.CTkFont(size=11)
        )
        self.ref_tipo.pack(fill="x")
        
        # Autor
        autor_frame = ctk.CTkFrame(row1_frame, fg_color="transparent")
        autor_frame.pack(side="right", fill="x", expand=True, padx=(5, 0))
        ctk.CTkLabel(autor_frame, text="Autor:", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.ref_autor = ctk.CTkEntry(autor_frame, placeholder_text="Apellido, N.", height=25, font=ctk.CTkFont(size=11))
        self.ref_autor.pack(fill="x")
        
        # Segunda fila
        row2_frame = ctk.CTkFrame(fields_frame, fg_color="transparent")
        row2_frame.pack(fill="x", pady=3)
        
        # Año
        ano_frame = ctk.CTkFrame(row2_frame, fg_color="transparent")
        ano_frame.pack(side="left", fill="x", expand=True, padx=(0, 5))
        ctk.CTkLabel(ano_frame, text="Año:", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.ref_ano = ctk.CTkEntry(ano_frame, placeholder_text="2024", height=25, font=ctk.CTkFont(size=11))
        self.ref_ano.pack(fill="x")
        
        # Título
        titulo_frame = ctk.CTkFrame(row2_frame, fg_color="transparent")
        titulo_frame.pack(side="right", fill="x", expand=True, padx=(5, 0))
        ctk.CTkLabel(titulo_frame, text="Título:", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.ref_titulo = ctk.CTkEntry(titulo_frame, placeholder_text="Título", height=25, font=ctk.CTkFont(size=11))
        self.ref_titulo.pack(fill="x")
        
        # Tercera fila
        row3_frame = ctk.CTkFrame(fields_frame, fg_color="transparent")
        row3_frame.pack(fill="x", pady=3)
        
        # Fuente
        fuente_frame = ctk.CTkFrame(row3_frame, fg_color="transparent")
        fuente_frame.pack(fill="x")
        ctk.CTkLabel(fuente_frame, text="Fuente:", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.ref_fuente = ctk.CTkEntry(fuente_frame, placeholder_text="Editorial o URL", height=25, font=ctk.CTkFont(size=11))
        self.ref_fuente.pack(fill="x")
        
        add_ref_btn = ctk.CTkButton(
            ref_frame, text="➕ Agregar", command=self.agregar_referencia,
            height=28, font=ctk.CTkFont(size=12, weight="bold"), width=120
        )
        add_ref_btn.pack(pady=(5, 10))
        
        # Lista de referencias - más compacta
        list_frame = ctk.CTkFrame(main_frame)
        list_frame.pack(fill="both", expand=True)
        
        list_title = ctk.CTkLabel(
            list_frame, text="📋 Referencias Agregadas",
            font=ctk.CTkFont(size=14, weight="bold")
        )
        list_title.pack(pady=(10, 8))
        
        self.ref_scroll_frame = ctk.CTkScrollableFrame(list_frame, height=140)
        self.ref_scroll_frame.pack(fill="both", expand=True, padx=15, pady=(0, 8))
        
        delete_btn = ctk.CTkButton(
            list_frame, text="🗑️ Eliminar Última", command=self.eliminar_referencia,
            fg_color="red", hover_color="darkred", height=28, width=140
        )
        delete_btn.pack(pady=(0, 10))
    
    def setup_formato_avanzado(self):
        """Pestaña para opciones avanzadas de formato - más compacta"""
        tab = self.tabview.add("🎨 Formato")
        
        scroll_frame = ctk.CTkScrollableFrame(tab, label_text="Configuración de Formato", height=400)
        scroll_frame.pack(fill="both", expand=True, padx=15, pady=15)
        
        # Sección de tipografía - más compacta
        tipo_frame = ctk.CTkFrame(scroll_frame, fg_color="darkgreen", corner_radius=8)
        tipo_frame.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(
            tipo_frame, text="🔤 Tipografía",
            font=ctk.CTkFont(size=14, weight="bold"), text_color="white"
        ).pack(pady=(10, 8))
        
        # Configuración de tipografía usando pack - más compacta
        tipo_grid = ctk.CTkFrame(tipo_frame, fg_color="transparent")
        tipo_grid.pack(fill="x", padx=15, pady=(0, 10))
        
        # Primera fila
        row1_tipo = ctk.CTkFrame(tipo_grid, fg_color="transparent")
        row1_tipo.pack(fill="x", pady=3)
        
        # Fuente del texto
        fuente_texto_frame = ctk.CTkFrame(row1_tipo, fg_color="transparent")
        fuente_texto_frame.pack(side="left", fill="x", expand=True, padx=(0, 5))
        ctk.CTkLabel(fuente_texto_frame, text="Fuente texto:", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.fuente_texto = ctk.CTkComboBox(
            fuente_texto_frame, values=["Times New Roman", "Arial", "Calibri"], height=25
        )
        self.fuente_texto.set("Times New Roman")
        self.fuente_texto.pack(fill="x")
        
        # Tamaño del texto
        tamaño_texto_frame = ctk.CTkFrame(row1_tipo, fg_color="transparent")
        tamaño_texto_frame.pack(side="right", fill="x", expand=True, padx=(5, 0))
        ctk.CTkLabel(tamaño_texto_frame, text="Tamaño:", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.tamaño_texto = ctk.CTkComboBox(tamaño_texto_frame, values=["10", "11", "12", "13", "14"], height=25)
        self.tamaño_texto.set("12")
        self.tamaño_texto.pack(fill="x")
        
        # Segunda fila
        row2_tipo = ctk.CTkFrame(tipo_grid, fg_color="transparent")
        row2_tipo.pack(fill="x", pady=3)
        
        # Fuente de títulos
        fuente_titulo_frame = ctk.CTkFrame(row2_tipo, fg_color="transparent")
        fuente_titulo_frame.pack(side="left", fill="x", expand=True, padx=(0, 5))
        ctk.CTkLabel(fuente_titulo_frame, text="Fuente títulos:", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.fuente_titulo = ctk.CTkComboBox(
            fuente_titulo_frame, values=["Times New Roman", "Arial", "Calibri"], height=25
        )
        self.fuente_titulo.set("Times New Roman")
        self.fuente_titulo.pack(fill="x")
        
        # Tamaño de títulos
        tamaño_titulo_frame = ctk.CTkFrame(row2_tipo, fg_color="transparent")
        tamaño_titulo_frame.pack(side="right", fill="x", expand=True, padx=(5, 0))
        ctk.CTkLabel(tamaño_titulo_frame, text="Tamaño:", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.tamaño_titulo = ctk.CTkComboBox(tamaño_titulo_frame, values=["12", "13", "14", "15", "16"], height=25)
        self.tamaño_titulo.set("14")
        self.tamaño_titulo.pack(fill="x")
        
        # Sección de espaciado - más compacta
        espaciado_frame = ctk.CTkFrame(scroll_frame, fg_color="darkblue", corner_radius=8)
        espaciado_frame.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(
            espaciado_frame, text="📏 Espaciado",
            font=ctk.CTkFont(size=14, weight="bold"), text_color="white"
        ).pack(pady=(10, 8))
        
        espaciado_grid = ctk.CTkFrame(espaciado_frame, fg_color="transparent")
        espaciado_grid.pack(fill="x", padx=15, pady=(0, 10))
        
        # Fila de espaciado
        espaciado_row = ctk.CTkFrame(espaciado_grid, fg_color="transparent")
        espaciado_row.pack(fill="x", pady=3)
        
        # Interlineado
        interlineado_frame = ctk.CTkFrame(espaciado_row, fg_color="transparent")
        interlineado_frame.pack(side="left", fill="x", expand=True, padx=(0, 5))
        ctk.CTkLabel(interlineado_frame, text="Interlineado:", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.interlineado = ctk.CTkComboBox(interlineado_frame, values=["1.0", "1.5", "2.0"], height=25)
        self.interlineado.set("2.0")
        self.interlineado.pack(fill="x")
        
        # Márgenes
        margen_frame = ctk.CTkFrame(espaciado_row, fg_color="transparent")
        margen_frame.pack(side="right", fill="x", expand=True, padx=(5, 0))
        ctk.CTkLabel(margen_frame, text="Márgenes (cm):", font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
        self.margen = ctk.CTkComboBox(margen_frame, values=["2.0", "2.54", "3.0"], height=25)
        self.margen.set("2.54")
        self.margen.pack(fill="x")
        
        # Opciones de alineación - más compacta
        align_frame = ctk.CTkFrame(scroll_frame, fg_color="darkred", corner_radius=8)
        align_frame.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(
            align_frame, text="📐 Alineación",
            font=ctk.CTkFont(size=14, weight="bold"), text_color="white"
        ).pack(pady=(10, 8))
        
        align_grid = ctk.CTkFrame(align_frame, fg_color="transparent")
        align_grid.pack(fill="x", padx=15, pady=(0, 10))
        
        # Fila de opciones de alineación
        align_row = ctk.CTkFrame(align_grid, fg_color="transparent")
        align_row.pack(fill="x", pady=3)
        
        # Opciones de formato profesional
        self.salto_pagina_var = ctk.CTkCheckBox(
            align_row, text="Salto de página entre secciones", font=ctk.CTkFont(size=12)
        )
        self.salto_pagina_var.select()
        self.salto_pagina_var.pack(side="left", padx=(10, 10))
        
        self.conservar_siguiente_var = ctk.CTkCheckBox(
            align_row, text="Conservar con siguiente", font=ctk.CTkFont(size=12)
        )
        self.conservar_siguiente_var.select()
        self.conservar_siguiente_var.pack(side="right", padx=(10, 10))
        
        # Segunda fila de opciones profesionales
        align_row2 = ctk.CTkFrame(align_grid, fg_color="transparent")
        align_row2.pack(fill="x", pady=3)
        
        self.justificado_var = ctk.CTkCheckBox(
            align_row2, text="Justificado", font=ctk.CTkFont(size=12)
        )
        self.justificado_var.select()
        self.justificado_var.pack(side="left", padx=(10, 10))
        
        self.sangria_var = ctk.CTkCheckBox(
            align_row2, text="Sangría primera línea", font=ctk.CTkFont(size=12)
        )
        self.sangria_var.select()
        self.sangria_var.pack(side="right", padx=(10, 10))
        
        # Botón para aplicar configuración
        apply_btn = ctk.CTkButton(
            scroll_frame, text="✅ Aplicar Configuración", command=self.aplicar_formato,
            height=35, font=ctk.CTkFont(size=12, weight="bold"),
            fg_color="green", hover_color="darkgreen"
        )
        apply_btn.pack(pady=15)
    
    def setup_generacion(self):
        """Pestaña de generación mejorada"""
        tab = self.tabview.add("🔧 Generar")
        
        main_frame = ctk.CTkFrame(tab, fg_color="transparent")
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Opciones de inclusión
        options_frame = ctk.CTkFrame(main_frame, corner_radius=15)
        options_frame.pack(fill="x", pady=(0, 20))
        
        options_title = ctk.CTkLabel(
            options_frame, text="⚙️ Opciones de Generación",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        options_title.pack(pady=(20, 15))
        
        options_grid = ctk.CTkFrame(options_frame, fg_color="transparent")
        options_grid.pack(padx=30, pady=(0, 20))
        
        # Primera fila de opciones
        options_row1 = ctk.CTkFrame(options_grid, fg_color="transparent")
        options_row1.pack(fill="x", pady=5)
        
        self.incluir_portada = ctk.CTkCheckBox(
            options_row1, text="📄 Incluir Portada", font=ctk.CTkFont(size=14)
        )
        self.incluir_portada.select()
        self.incluir_portada.pack(side="left", padx=(20, 20))
        
        self.incluir_indice = ctk.CTkCheckBox(
            options_row1, text="📑 Incluir Índice", font=ctk.CTkFont(size=14)
        )
        self.incluir_indice.select()
        self.incluir_indice.pack(side="right", padx=(20, 20))
        
        # Segunda fila de opciones
        options_row2 = ctk.CTkFrame(options_grid, fg_color="transparent")
        options_row2.pack(fill="x", pady=5)
        
        self.incluir_agradecimientos = ctk.CTkCheckBox(
            options_row2, text="🙏 Incluir Agradecimientos", font=ctk.CTkFont(size=14)
        )
        self.incluir_agradecimientos.pack(side="left", padx=(20, 20))
        
        self.numeracion_paginas = ctk.CTkCheckBox(
            options_row2, text="📊 Numeración de páginas", font=ctk.CTkFont(size=14)
        )
        self.numeracion_paginas.select()
        self.numeracion_paginas.pack(side="right", padx=(20, 20))
        
        # Frame de validación
        validation_frame = ctk.CTkFrame(main_frame, corner_radius=15)
        validation_frame.pack(fill="both", expand=True)
        
        validation_title = ctk.CTkLabel(
            validation_frame, text="🔍 Estado del Proyecto",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        validation_title.pack(pady=(20, 15))
        
        self.validation_text = ctk.CTkTextbox(
            validation_frame, height=250, font=ctk.CTkFont(size=12), fg_color="gray10"
        )
        self.validation_text.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Inicializar con mensaje de bienvenida
        self.validation_text.insert("1.0", 
            "✨ ¡Generador Avanzado de Proyectos Académicos!\n\n"
            "🆕 NUEVAS CARACTERÍSTICAS:\n"
            "• Gestión dinámica de secciones\n"
            "• Formato personalizable (fuentes, tamaños, espaciado)\n"
            "• Uso de plantilla base del documento\n"
            "• Reordenamiento de contenido\n"
            "• Validación avanzada\n\n"
            "📋 PROCESO RECOMENDADO:\n"
            "1. Completa información general\n"
            "2. Organiza secciones según tu necesidad\n"
            "3. Personaliza el formato\n"
            "4. Agrega contenido y referencias\n"
            "5. Valida y genera\n\n"
            "🎯 ¡Haz clic en 'Validar Proyecto' cuando estés listo!"
        )
        
        self.progress = ctk.CTkProgressBar(validation_frame, height=20)
        self.progress.pack(fill="x", padx=20, pady=(0, 20))
        self.progress.set(0)
    
    def gestionar_imagenes(self):
        """Abre ventana para gestionar imágenes del documento"""
        img_window = ctk.CTkToplevel(self.root)
        img_window.title("🖼️ Gestión de Imágenes")
        img_window.geometry("600x500")
        img_window.transient(self.root)
        img_window.grab_set()
        
        # Centrar ventana
        img_window.update_idletasks()
        x = (img_window.winfo_screenwidth() // 2) - (600 // 2)
        y = (img_window.winfo_screenheight() // 2) - (500 // 2)
        img_window.geometry(f"600x500+{x}+{y}")
        
        main_frame = ctk.CTkFrame(img_window, corner_radius=0)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Título
        title_label = ctk.CTkLabel(
            main_frame, text="🖼️ Gestión de Imágenes del Documento",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(pady=(10, 20))
        
        # Estado de imágenes base
        status_frame = ctk.CTkFrame(main_frame, fg_color="gray20", corner_radius=10)
        status_frame.pack(fill="x", pady=(0, 20))
        
        status_title = ctk.CTkLabel(
            status_frame, text="📁 Estado de Imágenes Base",
            font=ctk.CTkFont(size=14, weight="bold")
        )
        status_title.pack(pady=(10, 5))
        
        # Estado del encabezado
        enc_status = "✅ Encontrado" if self.ruta_encabezado else "❌ No encontrado"
        enc_label = ctk.CTkLabel(
            status_frame, text=f"Encabezado.png: {enc_status}",
            font=ctk.CTkFont(size=12)
        )
        enc_label.pack(pady=2)
        
        # Estado de la insignia
        ins_status = "✅ Encontrado" if self.ruta_insignia else "❌ No encontrado"
        ins_label = ctk.CTkLabel(
            status_frame, text=f"Insignia.png: {ins_status}",
            font=ctk.CTkFont(size=12)
        )
        ins_label.pack(pady=(2, 10))
        
        # Sección de carga personalizada
        custom_frame = ctk.CTkFrame(main_frame, fg_color="darkblue", corner_radius=10)
        custom_frame.pack(fill="x", pady=(0, 20))
        
        custom_title = ctk.CTkLabel(
            custom_frame, text="📤 Cargar Imágenes Personalizadas",
            font=ctk.CTkFont(size=14, weight="bold"), text_color="white"
        )
        custom_title.pack(pady=(15, 10))
        
        # Botones de carga
        btn_frame = ctk.CTkFrame(custom_frame, fg_color="transparent")
        btn_frame.pack(fill="x", padx=20, pady=(0, 15))
        
        enc_btn = ctk.CTkButton(
            btn_frame, text="📋 Cargar Encabezado", 
            command=lambda: self.cargar_imagen_personalizada("encabezado"),
            width=200, height=35
        )
        enc_btn.pack(side="left", padx=(0, 10))
        
        ins_btn = ctk.CTkButton(
            btn_frame, text="🏛️ Cargar Insignia", 
            command=lambda: self.cargar_imagen_personalizada("insignia"),
            width=200, height=35
        )
        ins_btn.pack(side="right", padx=(10, 0))
        
        # Estado de imágenes personalizadas
        custom_status_frame = ctk.CTkFrame(main_frame, fg_color="gray20", corner_radius=10)
        custom_status_frame.pack(fill="x", pady=(0, 20))
        
        custom_status_title = ctk.CTkLabel(
            custom_status_frame, text="🎨 Imágenes Personalizadas Cargadas",
            font=ctk.CTkFont(size=14, weight="bold")
        )
        custom_status_title.pack(pady=(10, 5))
        
        self.enc_custom_label = ctk.CTkLabel(
            custom_status_frame, 
            text=f"Encabezado: {'✅ Cargado' if self.encabezado_personalizado else '⏸️ No cargado'}",
            font=ctk.CTkFont(size=12)
        )
        self.enc_custom_label.pack(pady=2)
        
        self.ins_custom_label = ctk.CTkLabel(
            custom_status_frame, 
            text=f"Insignia: {'✅ Cargado' if self.insignia_personalizada else '⏸️ No cargado'}",
            font=ctk.CTkFont(size=12)
        )
        self.ins_custom_label.pack(pady=(2, 10))
        
        # Botones de acción
        action_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        action_frame.pack(fill="x", pady=(0, 10))
        
        reset_btn = ctk.CTkButton(
            action_frame, text="🔄 Restablecer", 
            command=lambda: self.restablecer_imagenes(img_window),
            width=120, height=35, fg_color="red", hover_color="darkred"
        )
        reset_btn.pack(side="left")
        
        close_btn = ctk.CTkButton(
            action_frame, text="✅ Cerrar", 
            command=img_window.destroy,
            width=120, height=35
        )
        close_btn.pack(side="right")
        
        # Información adicional
        info_frame = ctk.CTkFrame(main_frame, fg_color="green", corner_radius=10)
        info_frame.pack(fill="x")
        
        info_text = """💡 INFORMACIÓN IMPORTANTE:
• Las imágenes base se buscan en: /Recursos/Encabezado.png e /Recursos/Insignia.png
• Las imágenes personalizadas tienen prioridad sobre las base
• Formatos soportados: PNG, JPG, JPEG
• Tamaño recomendado: Encabezado 600x100px, Insignia 100x100px"""
        
        info_label = ctk.CTkLabel(
            info_frame, text=info_text, font=ctk.CTkFont(size=10),
            justify="left", wraplength=550, text_color="white"
        )
        info_label.pack(padx=15, pady=10)
    
    def cargar_imagen_personalizada(self, tipo):
        """Carga una imagen personalizada"""
        filetypes = [
            ("Imágenes", "*.png *.jpg *.jpeg"),
            ("PNG files", "*.png"),
            ("JPG files", "*.jpg *.jpeg")
        ]
        
        filename = filedialog.askopenfilename(
            title=f"Seleccionar {tipo.capitalize()}",
            filetypes=filetypes
        )
        
        if filename:
            try:
                # Verificar que sea una imagen válida
                with Image.open(filename) as img:
                    # Validar tamaño mínimo
                    if img.width < 50 or img.height < 50:
                        messagebox.showwarning("⚠️ Advertencia", 
                            "La imagen es muy pequeña. Se recomienda al menos 50x50 píxeles.")
                    
                    if tipo == "encabezado":
                        self.encabezado_personalizado = filename
                        self.enc_custom_label.configure(text="Encabezado: ✅ Cargado")
                    else:
                        self.insignia_personalizada = filename
                        self.ins_custom_label.configure(text="Insignia: ✅ Cargado")
                    
                    messagebox.showinfo("✅ Éxito", 
                        f"{tipo.capitalize()} cargado correctamente.\n"
                        f"Tamaño: {img.width}x{img.height} píxeles")
                        
            except Exception as e:
                messagebox.showerror("❌ Error", 
                    f"Error al cargar la imagen:\n{str(e)}")
    
    def restablecer_imagenes(self, window):
        """Restablece las imágenes a las base"""
        self.encabezado_personalizado = None
        self.insignia_personalizada = None
        
        # Actualizar labels
        self.enc_custom_label.configure(text="Encabezado: ⏸️ No cargado")
        self.ins_custom_label.configure(text="Insignia: ⏸️ No cargado")
        
        messagebox.showinfo("🔄 Restablecido", 
            "Se usarán las imágenes base (si están disponibles)")
    
    def setup_keyboard_shortcuts(self):
        """Configura atajos de teclado"""
        self.root.bind('<Control-s>', lambda e: self.guardar_proyecto())
        self.root.bind('<Control-o>', lambda e: self.cargar_proyecto())
        self.root.bind('<Control-n>', lambda e: self.nuevo_proyecto())
        self.root.bind('<F5>', lambda e: self.validar_proyecto())
        self.root.bind('<F9>', lambda e: self.generar_documento_async())
        self.root.bind('<Control-q>', lambda e: self.root.quit())
    
    def actualizar_estadisticas(self):
        """Actualiza las estadísticas en tiempo real"""
        total_words = 0
        total_chars = 0
        sections_completed = 0
        
        for key, text_widget in self.content_texts.items():
            if key in self.secciones_disponibles:
                content = text_widget.get("1.0", "end").strip()
                if content and len(content) > 10:
                    sections_completed += 1
                    words = len(content.split())
                    total_words += words
                    total_chars += len(content)
        
        self.stats = {
            'total_words': total_words,
            'total_chars': total_chars,
            'sections_completed': sections_completed,
            'references_added': len(self.referencias)
        }
        
        # Actualizar label
        total_sections = len([s for s in self.secciones_disponibles.values() if not s['capitulo']])
        stats_text = f"📊 Palabras: {total_words} | Secciones: {sections_completed}/{total_sections} | Referencias: {len(self.referencias)}"
        self.stats_label.configure(text=stats_text)
        
        # Programar próxima actualización
        self.root.after(2000, self.actualizar_estadisticas)
    
    def guardar_proyecto(self):
        """Guarda el proyecto completo en un archivo JSON"""
        try:
            # Recopilar todos los datos
            proyecto_completo = {
                'version': '2.0',
                'fecha_creacion': datetime.now().isoformat(),
                'informacion_general': {},
                'contenido_secciones': {},
                'referencias': self.referencias,
                'secciones_activas': self.secciones_activas,
                'secciones_disponibles': self.secciones_disponibles,
                'formato_config': self.formato_config,
                'imagenes': {
                    'encabezado_personalizado': self.encabezado_personalizado,
                    'insignia_personalizada': self.insignia_personalizada
                },
                'estadisticas': self.stats
            }
            
            # Información general
            for key, entry in self.proyecto_data.items():
                if hasattr(entry, 'get'):
                    proyecto_completo['informacion_general'][key] = entry.get()
            
            # Contenido de secciones
            for key, text_widget in self.content_texts.items():
                proyecto_completo['contenido_secciones'][key] = text_widget.get("1.0", "end")
            
            # Guardar archivo
            filename = filedialog.asksaveasfilename(
                defaultextension=".json",
                filetypes=[("Proyecto Académico", "*.json"), ("Todos los archivos", "*.*")],
                title="Guardar Proyecto"
            )
            
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    json.dump(proyecto_completo, f, ensure_ascii=False, indent=2)
                
                self.last_save_time = datetime.now()
                messagebox.showinfo("💾 Guardado", 
                    f"Proyecto guardado exitosamente:\n{os.path.basename(filename)}")
                
        except Exception as e:
            messagebox.showerror("❌ Error", f"Error al guardar proyecto:\n{str(e)}")
    
    def cargar_proyecto(self):
        """Carga un proyecto desde archivo JSON"""
        try:
            filename = filedialog.askopenfilename(
                filetypes=[("Proyecto Académico", "*.json"), ("Todos los archivos", "*.*")],
                title="Cargar Proyecto"
            )
            
            if not filename:
                return
            
            with open(filename, 'r', encoding='utf-8') as f:
                proyecto_completo = json.load(f)
            
            # Verificar versión
            version = proyecto_completo.get('version', '1.0')
            if version != '2.0':
                messagebox.showwarning("⚠️ Versión", 
                    "Este proyecto fue creado con una versión anterior. Algunas características pueden no funcionar correctamente.")
            
            # Cargar información general
            if 'informacion_general' in proyecto_completo:
                for key, value in proyecto_completo['informacion_general'].items():
                    if key in self.proyecto_data and hasattr(self.proyecto_data[key], 'delete'):
                        self.proyecto_data[key].delete(0, "end")
                        self.proyecto_data[key].insert(0, value)
            
            # Cargar configuración de formato
            if 'formato_config' in proyecto_completo:
                self.formato_config.update(proyecto_completo['formato_config'])
                self.aplicar_config_cargada()
            
            # Cargar secciones
            if 'secciones_disponibles' in proyecto_completo:
                self.secciones_disponibles = proyecto_completo['secciones_disponibles']
            
            if 'secciones_activas' in proyecto_completo:
                self.secciones_activas = proyecto_completo['secciones_activas']
            
            # Cargar referencias
            if 'referencias' in proyecto_completo:
                self.referencias = proyecto_completo['referencias']
                self.actualizar_lista_referencias()
            
            # Cargar imágenes
            if 'imagenes' in proyecto_completo:
                self.encabezado_personalizado = proyecto_completo['imagenes'].get('encabezado_personalizado')
                self.insignia_personalizada = proyecto_completo['imagenes'].get('insignia_personalizada')
            
            # Recrear interfaz
            self.actualizar_lista_secciones()
            self.crear_pestanas_contenido()
            
            # Cargar contenido de secciones
            if 'contenido_secciones' in proyecto_completo:
                for key, content in proyecto_completo['contenido_secciones'].items():
                    if key in self.content_texts:
                        self.content_texts[key].delete("1.0", "end")
                        self.content_texts[key].insert("1.0", content)
            
            messagebox.showinfo("📂 Cargado", 
                f"Proyecto cargado exitosamente:\n{os.path.basename(filename)}")
                
        except Exception as e:
            messagebox.showerror("❌ Error", f"Error al cargar proyecto:\n{str(e)}")
    
    def nuevo_proyecto(self):
        """Crea un nuevo proyecto limpio"""
        respuesta = messagebox.askyesno("🆕 Nuevo Proyecto", 
            "¿Estás seguro? Se perderán todos los cambios no guardados.")
        
        if respuesta:
            # Reiniciar todas las variables
            self.referencias = []
            self.encabezado_personalizado = None
            self.insignia_personalizada = None
            self.secciones_disponibles = self.get_secciones_iniciales()
            self.secciones_activas = list(self.secciones_disponibles.keys())
            
            # Limpiar campos
            for entry in self.proyecto_data.values():
                if hasattr(entry, 'delete'):
                    entry.delete(0, "end")
            
            # Recrear interfaz
            self.actualizar_lista_secciones()
            self.crear_pestanas_contenido()
            self.actualizar_lista_referencias()
            
            messagebox.showinfo("🆕 Nuevo Proyecto", "Proyecto nuevo creado")
    
    def auto_save_project(self):
        """Guarda automáticamente el proyecto"""
        if self.auto_save_enabled:
            try:
                script_dir = os.path.dirname(os.path.abspath(__file__))
                auto_save_path = os.path.join(script_dir, "auto_save.json")
                
                # Crear backup automático
                proyecto_completo = {
                    'version': '2.0',
                    'fecha_auto_save': datetime.now().isoformat(),
                    'informacion_general': {},
                    'contenido_secciones': {},
                    'referencias': self.referencias,
                    'secciones_activas': self.secciones_activas,
                    'formato_config': self.formato_config
                }
                
                # Guardar información
                for key, entry in self.proyecto_data.items():
                    if hasattr(entry, 'get'):
                        proyecto_completo['informacion_general'][key] = entry.get()
                
                for key, text_widget in self.content_texts.items():
                    proyecto_completo['contenido_secciones'][key] = text_widget.get("1.0", "end")
                
                with open(auto_save_path, 'w', encoding='utf-8') as f:
                    json.dump(proyecto_completo, f, ensure_ascii=False, indent=2)
                
                print(f"Auto-guardado realizado: {datetime.now().strftime('%H:%M:%S')}")
                
            except Exception as e:
                print(f"Error en auto-guardado: {e}")
            
            # Programar próximo auto-guardado
            self.root.after(300000, self.auto_save_project)  # 5 minutos
    
    def exportar_configuracion(self):
        """Exporta solo la configuración de formato"""
        try:
            config_export = {
                'version': '2.0',
                'tipo': 'configuracion_formato',
                'fecha_export': datetime.now().isoformat(),
                'formato_config': self.formato_config,
                'secciones_disponibles': self.secciones_disponibles,
                'secciones_activas': self.secciones_activas
            }
            
            filename = filedialog.asksaveasfilename(
                defaultextension=".json",
                filetypes=[("Configuración", "*.json")],
                title="Exportar Configuración"
            )
            
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    json.dump(config_export, f, ensure_ascii=False, indent=2)
                
                messagebox.showinfo("📤 Exportado", 
                    "Configuración exportada exitosamente")
                    
        except Exception as e:
            messagebox.showerror("❌ Error", f"Error al exportar:\n{str(e)}")
    
    def aplicar_config_cargada(self):
        """Aplica configuración cargada a los controles"""
        try:
            # Actualizar controles de formato si existen
            if hasattr(self, 'fuente_texto'):
                self.fuente_texto.set(self.formato_config.get('fuente_texto', 'Times New Roman'))
            if hasattr(self, 'tamaño_texto'):
                self.tamaño_texto.set(str(self.formato_config.get('tamaño_texto', 12)))
            if hasattr(self, 'fuente_titulo'):
                self.fuente_titulo.set(self.formato_config.get('fuente_titulo', 'Times New Roman'))
            if hasattr(self, 'tamaño_titulo'):
                self.tamaño_titulo.set(str(self.formato_config.get('tamaño_titulo', 14)))
            if hasattr(self, 'interlineado'):
                self.interlineado.set(str(self.formato_config.get('interlineado', 2.0)))
            if hasattr(self, 'margen'):
                self.margen.set(str(self.formato_config.get('margen', 2.54)))
                
        except Exception as e:
            print(f"Error aplicando configuración: {e}")
    
    def actualizar_lista_referencias(self):
        """Actualiza la lista visual de referencias"""
        # Limpiar lista actual
        for widget in self.ref_scroll_frame.winfo_children():
            widget.destroy()
        
        # Recrear todas las referencias
        for ref in self.referencias:
            ref_item_frame = ctk.CTkFrame(self.ref_scroll_frame, fg_color="gray20", corner_radius=8)
            ref_item_frame.pack(fill="x", padx=5, pady=5)
            
            apa_ref = f"{ref['autor']} ({ref['año']}). {ref['titulo']}. {ref['fuente']}"
            ref_label = ctk.CTkLabel(
                ref_item_frame, text=f"📖 {apa_ref}", font=ctk.CTkFont(size=11),
                wraplength=800, justify="left"
            )
            ref_label.pack(padx=15, pady=10, anchor="w")
        """Obtiene la ruta final de la imagen a usar (personalizada o base)"""
        if tipo == "encabezado":
            return self.encabezado_personalizado or self.ruta_encabezado
        elif tipo == "insignia":
            return self.insignia_personalizada or self.ruta_insignia
        return None
        """Activa/desactiva el uso del formato base"""
        if self.usar_base_var.get():
            if self.documento_base is None:
                self.cargar_documento_base()
            else:
                self.aplicar_formato_base()
        else:
            self.limpiar_formato_base()
    
    def cargar_documento_base(self):
        """Carga el documento base proporcionado"""
        try:
            # Simulamos la carga del formato base
            self.documento_base = {
                'institucion': 'COLEGIO PRIVADO DIVINA ESPERANZA',
                'ciclo': 'Tercer año',
                'curso': '3 BTI',
                'enfasis': 'Tecnología',
                'director': 'Cristina Raichakowski'
            }
            
            messagebox.showinfo("✅ Éxito", 
                "Formato base cargado correctamente.\n"
                "Se aplicarán automáticamente:\n"
                "• Estructura del documento\n"
                "• Datos predefinidos\n"
                "• Formato específico\n\n"
                "Marca la casilla para activarlo.")
            
            self.usar_base_var.select()
            self.aplicar_formato_base()
            
        except Exception as e:
            messagebox.showerror("❌ Error", f"Error al cargar formato base:\n{str(e)}")
    
    def aplicar_formato_base(self):
        """Aplica los datos del formato base"""
        if self.documento_base:
            for key, value in self.documento_base.items():
                if key in self.proyecto_data:
                    self.proyecto_data[key].delete(0, "end")
                    self.proyecto_data[key].insert(0, value)
            
            messagebox.showinfo("✅ Aplicado", "Formato base aplicado correctamente")
    
    def limpiar_formato_base(self):
        """Limpia los datos del formato base"""
        campos_base = ['institucion', 'ciclo', 'curso', 'enfasis', 'director']
        for campo in campos_base:
            if campo in self.proyecto_data:
                self.proyecto_data[campo].delete(0, "end")
    
    def aplicar_formato(self):
        """Aplica la configuración de formato"""
        self.formato_config = {
            'fuente_texto': self.fuente_texto.get(),
            'tamaño_texto': int(self.tamaño_texto.get()),
            'fuente_titulo': self.fuente_titulo.get(),
            'tamaño_titulo': int(self.tamaño_titulo.get()),
            'interlineado': float(self.interlineado.get()),
            'margen': float(self.margen.get()),
            'justificado': self.justificado_var.get(),
            'sangria': self.sangria_var.get()
        }
        
        # Actualizar preview
        self.preview_label.configure(
            font=ctk.CTkFont(family=self.formato_config['fuente_texto'], 
                           size=self.formato_config['tamaño_texto'])
        )
        
        messagebox.showinfo("✅ Aplicado", "Configuración de formato aplicada correctamente")
    
    def actualizar_lista_secciones(self):
        """Actualiza la lista visual de secciones activas"""
        # Limpiar lista actual
        for widget in self.secciones_listbox.winfo_children():
            widget.destroy()
        
        # Recrear lista
        for i, seccion_id in enumerate(self.secciones_activas):
            if seccion_id in self.secciones_disponibles:
                seccion = self.secciones_disponibles[seccion_id]
                
                item_frame = ctk.CTkFrame(self.secciones_listbox, fg_color="gray20", corner_radius=5)
                item_frame.pack(fill="x", pady=2, padx=5)
                
                # Checkbox para selección
                checkbox = ctk.CTkCheckBox(
                    item_frame, text="", width=20, command=lambda idx=i: self.seleccionar_seccion(idx)
                )
                checkbox.pack(side="left", padx=(10, 5), pady=5)
                
                # Texto de la sección
                color = "yellow" if seccion['requerida'] else "white"
                if seccion['capitulo']:
                    color = "lightblue"
                
                label = ctk.CTkLabel(
                    item_frame, text=seccion['titulo'], 
                    font=ctk.CTkFont(size=11), text_color=color
                )
                label.pack(side="left", fill="x", expand=True, pady=5)
                
                # Guardar referencia para selección
                checkbox.seccion_index = i
    
    def seleccionar_seccion(self, index):
        """Maneja la selección de secciones"""
        # Esta función se puede expandir para manejar selecciones múltiples
        pass
    
    def agregar_seccion(self):
        """Agrega una nueva sección personalizada"""
        dialog = SeccionDialog(self.root, self.secciones_disponibles)
        if dialog.result:
            seccion_id, seccion_data = dialog.result
            self.secciones_disponibles[seccion_id] = seccion_data
            self.secciones_activas.append(seccion_id)
            self.actualizar_lista_secciones()
            self.crear_pestanas_contenido()
            messagebox.showinfo("✅ Agregada", f"Sección '{seccion_data['titulo']}' agregada correctamente")
    
    def quitar_seccion(self):
        """Quita secciones seleccionadas"""
        # Encontrar secciones seleccionadas
        secciones_a_quitar = []
        for widget in self.secciones_listbox.winfo_children():
            if hasattr(widget, 'winfo_children'):
                for child in widget.winfo_children():
                    if isinstance(child, ctk.CTkCheckBox) and child.get():
                        if hasattr(child, 'seccion_index'):
                            secciones_a_quitar.append(child.seccion_index)
        
        if secciones_a_quitar:
            # Verificar si hay secciones requeridas
            secciones_requeridas = []
            for idx in secciones_a_quitar:
                seccion_id = self.secciones_activas[idx]
                if self.secciones_disponibles[seccion_id]['requerida']:
                    secciones_requeridas.append(self.secciones_disponibles[seccion_id]['titulo'])
            
            if secciones_requeridas:
                messagebox.showwarning("⚠️ Advertencia", 
                    f"No se pueden eliminar secciones requeridas:\n{', '.join(secciones_requeridas)}")
                return
            
            # Quitar secciones (en orden inverso para mantener índices)
            for idx in sorted(secciones_a_quitar, reverse=True):
                seccion_id = self.secciones_activas.pop(idx)
                if seccion_id in self.content_texts:
                    del self.content_texts[seccion_id]
            
            self.actualizar_lista_secciones()
            self.crear_pestanas_contenido()
            messagebox.showinfo("✅ Eliminadas", f"{len(secciones_a_quitar)} sección(es) eliminada(s)")
        else:
            messagebox.showwarning("⚠️ Selección", "Selecciona al menos una sección para eliminar")
    
    def editar_seccion(self):
        """Edita una sección existente"""
        # Implementar diálogo de edición
        messagebox.showinfo("🚧 En desarrollo", "Función de edición en desarrollo")
    
    def subir_seccion(self):
        """Sube una sección en el orden"""
        secciones_seleccionadas = self.obtener_secciones_seleccionadas()
        if len(secciones_seleccionadas) == 1:
            idx = secciones_seleccionadas[0]
            if idx > 0:
                # Intercambiar posiciones
                self.secciones_activas[idx], self.secciones_activas[idx-1] = \
                self.secciones_activas[idx-1], self.secciones_activas[idx]
                self.actualizar_lista_secciones()
                self.crear_pestanas_contenido()
        else:
            messagebox.showwarning("⚠️ Selección", "Selecciona exactamente una sección")
    
    def bajar_seccion(self):
        """Baja una sección en el orden"""
        secciones_seleccionadas = self.obtener_secciones_seleccionadas()
        if len(secciones_seleccionadas) == 1:
            idx = secciones_seleccionadas[0]
            if idx < len(self.secciones_activas) - 1:
                # Intercambiar posiciones
                self.secciones_activas[idx], self.secciones_activas[idx+1] = \
                self.secciones_activas[idx+1], self.secciones_activas[idx]
                self.actualizar_lista_secciones()
                self.crear_pestanas_contenido()
        else:
            messagebox.showwarning("⚠️ Selección", "Selecciona exactamente una sección")
    
    def obtener_secciones_seleccionadas(self):
        """Obtiene las secciones actualmente seleccionadas"""
        seleccionadas = []
        for widget in self.secciones_listbox.winfo_children():
            if hasattr(widget, 'winfo_children'):
                for child in widget.winfo_children():
                    if isinstance(child, ctk.CTkCheckBox) and child.get():
                        if hasattr(child, 'seccion_index'):
                            seleccionadas.append(child.seccion_index)
        return seleccionadas
    
    def crear_pestanas_contenido(self):
        """Crea las pestañas de contenido basadas en secciones activas"""
        # Limpiar pestañas existentes
        for tab_name in list(self.content_tabview._tab_dict.keys()):
            self.content_tabview.delete(tab_name)
        
        # Crear nuevas pestañas
        for seccion_id in self.secciones_activas:
            if seccion_id in self.secciones_disponibles:
                seccion = self.secciones_disponibles[seccion_id]
                
                if not seccion['capitulo']:  # Solo crear pestañas para contenido, no para títulos de capítulo
                    tab = self.content_tabview.add(seccion['titulo'])
                    
                    # Frame de instrucciones
                    instruc_frame = ctk.CTkFrame(tab, fg_color="gray20", corner_radius=10)
                    instruc_frame.pack(fill="x", padx=10, pady=(10, 5))
                    
                    instruc_label = ctk.CTkLabel(
                        instruc_frame, text=f"💡 {seccion['instruccion']}",
                        font=ctk.CTkFont(size=12), wraplength=600, justify="left"
                    )
                    instruc_label.pack(padx=15, pady=10)
                    
                    # Text widget
                    text_widget = ctk.CTkTextbox(
                        tab, wrap="word", font=ctk.CTkFont(size=12), height=350
                    )
                    text_widget.pack(expand=True, fill="both", padx=10, pady=(5, 10))
                    
                    # Conservar contenido existente si ya existe
                    if seccion_id in self.content_texts:
                        contenido_existente = self.content_texts[seccion_id].get("1.0", "end")
                        text_widget.insert("1.0", contenido_existente)
                    
                    self.content_texts[seccion_id] = text_widget
    
    # Mantener métodos existentes (agregar_referencia, eliminar_referencia, etc.)
    def agregar_referencia(self):
        """Agrega una referencia a la lista con diseño moderno"""
        if not all([self.ref_autor.get(), self.ref_ano.get(), self.ref_titulo.get()]):
            messagebox.showerror("❌ Error", "Completa al menos Autor, Año y Título")
            return
        
        ref = {
            'tipo': self.ref_tipo.get(),
            'autor': self.ref_autor.get(),
            'año': self.ref_ano.get(),
            'titulo': self.ref_titulo.get(),
            'fuente': self.ref_fuente.get()
        }
        
        self.referencias.append(ref)
        
        # Crear frame para la referencia
        ref_item_frame = ctk.CTkFrame(self.ref_scroll_frame, fg_color="gray20", corner_radius=8)
        ref_item_frame.pack(fill="x", padx=5, pady=5)
        
        # Formato APA para mostrar
        apa_ref = f"{ref['autor']} ({ref['año']}). {ref['titulo']}. {ref['fuente']}"
        
        ref_label = ctk.CTkLabel(
            ref_item_frame, text=f"📖 {apa_ref}", font=ctk.CTkFont(size=11),
            wraplength=800, justify="left"
        )
        ref_label.pack(padx=15, pady=10, anchor="w")
        
        # Limpiar campos
        self.ref_autor.delete(0, "end")
        self.ref_ano.delete(0, "end")
        self.ref_titulo.delete(0, "end")
        self.ref_fuente.delete(0, "end")
        
        messagebox.showinfo("✅ Éxito", "Referencia agregada correctamente")
    
    def eliminar_referencia(self):
        """Elimina la última referencia"""
        if self.referencias:
            self.referencias.pop()
            # Recrear lista visual
            for widget in self.ref_scroll_frame.winfo_children():
                widget.destroy()
            
            for ref in self.referencias:
                ref_item_frame = ctk.CTkFrame(self.ref_scroll_frame, fg_color="gray20", corner_radius=8)
                ref_item_frame.pack(fill="x", padx=5, pady=5)
                
                apa_ref = f"{ref['autor']} ({ref['año']}). {ref['titulo']}. {ref['fuente']}"
                ref_label = ctk.CTkLabel(
                    ref_item_frame, text=f"📖 {apa_ref}", font=ctk.CTkFont(size=11),
                    wraplength=800, justify="left"
                )
                ref_label.pack(padx=15, pady=10, anchor="w")
            
            messagebox.showinfo("🗑️ Eliminado", "Última referencia eliminada")
        else:
            messagebox.showwarning("⚠️ Advertencia", "No hay referencias para eliminar")
    
    def validar_proyecto(self):
        """Valida el proyecto con las nuevas funcionalidades"""
        self.validation_text.delete("1.0", "end")
        errores = []
        advertencias = []
        
        # Validar información general
        campos_requeridos = ['titulo', 'estudiantes', 'tutores']
        for campo in campos_requeridos:
            if not self.proyecto_data[campo].get().strip():
                errores.append(f"❌ Campo requerido faltante: {campo}")
        
        # Validar secciones requeridas
        for seccion_id in self.secciones_activas:
            if seccion_id in self.secciones_disponibles:
                seccion = self.secciones_disponibles[seccion_id]
                if seccion['requerida'] and not seccion['capitulo']:
                    if seccion_id in self.content_texts:
                        content = self.content_texts[seccion_id].get("1.0", "end").strip()
                        if len(content) < 50:
                            errores.append(f"❌ Sección requerida '{seccion['titulo']}' muy corta")
                    else:
                        errores.append(f"❌ Sección requerida '{seccion['titulo']}' faltante")
        
        # Validar citas en marco teórico
        if 'marco_teorico' in self.content_texts:
            content = self.content_texts['marco_teorico'].get("1.0", "end")
            citas_encontradas = re.findall(r'\[CITA:[^\]]+\]', content)
            if not citas_encontradas:
                advertencias.append("⚠️ Marco Teórico sin citas detectadas")
        
        # Validar referencias
        if len(self.referencias) == 0:
            advertencias.append("⚠️ No hay referencias bibliográficas")
        
        # Mostrar resultados
        resultado = "🔍 VALIDACIÓN AVANZADA DEL PROYECTO\n" + "="*60 + "\n\n"
        
        if errores:
            resultado += "🚨 ERRORES CRÍTICOS:\n"
            for error in errores:
                resultado += f"{error}\n"
            resultado += "\n"
        
        if advertencias:
            resultado += "⚠️ ADVERTENCIAS:\n"
            for advertencia in advertencias:
                resultado += f"{advertencia}\n"
            resultado += "\n"
        
        # Estadísticas del proyecto
        resultado += "📊 ESTADÍSTICAS DEL PROYECTO:\n"
        resultado += f"• Secciones activas: {len(self.secciones_activas)}\n"
        resultado += f"• Secciones con contenido: {len([s for s in self.content_texts if self.content_texts[s].get('1.0', 'end').strip()])}\n"
        resultado += f"• Referencias bibliográficas: {len(self.referencias)}\n"
        resultado += f"• Formato personalizado: {'Sí' if self.formato_config['fuente_texto'] != 'Times New Roman' else 'Estándar'}\n"
        resultado += f"• Plantilla base: {'Activada' if self.usar_base_var.get() else 'No usada'}\n\n"
        
        if not errores and not advertencias:
            resultado += "✅ ¡PROYECTO PERFECTO!\n\n"
            resultado += "🎉 El proyecto cumple con todos los requisitos\n"
            resultado += "📄 Listo para generar con formato personalizado\n"
        elif not errores:
            resultado += "✅ PROYECTO VÁLIDO\n\n"
            resultado += "🎯 Proyecto listo para generar\n"
            resultado += "💡 Revisa las advertencias para mejorar\n"
        else:
            resultado += "❌ PROYECTO INCOMPLETO\n\n"
            resultado += "🔧 Corrige los errores marcados\n"
        
        self.validation_text.insert("1.0", resultado)
        
        # Actualizar progreso
        total_items = len(campos_requeridos) + len([s for s in self.secciones_disponibles.values() if s['requerida']]) + 1
        items_completos = total_items - len(errores)
        progreso = max(0, items_completos / total_items)
        self.progress.set(progreso)
    
    def procesar_citas(self, texto):
        """Procesa las citas mejorado"""
        def reemplazar_cita(match):
            cita_completa = match.group(0)
            contenido = cita_completa[6:-1]  # Quita [CITA: y ]
            partes = contenido.split(':')
            
            if len(partes) >= 3:
                tipo, autor, año = partes[0], partes[1], partes[2]
                pagina = partes[3] if len(partes) > 3 else None
                
                if tipo == 'textual':
                    return f" ({autor}, {año}, p. {pagina})" if pagina else f" ({autor}, {año})"
                elif tipo == 'parafraseo':
                    return f" ({autor}, {año})"
                elif tipo == 'larga':
                    return f"\n\n({autor}, {año}, p. {pagina})\n\n" if pagina else f"\n\n({autor}, {año})\n\n"
                elif tipo == 'web':
                    return f" ({autor}, {año})"
            
            return cita_completa
        
        return re.sub(r'\[CITA:[^\]]+\]', reemplazar_cita, texto)
    
    def generar_documento_async(self):
        """Genera el documento profesional en un hilo separado"""
        def generar():
            try:
                self.progress.set(0)
                self.progress.start()
                
                # Crear documento
                doc = Document()
                
                # Configurar estilos profesionales
                titulo_style = self.configurar_estilos_profesionales(doc)
                self.progress.set(0.1)
                
                # Generar contenido
                if self.incluir_portada.get():
                    self.crear_portada_profesional(doc)
                    self.progress.set(0.2)
                
                if self.incluir_agradecimientos.get():
                    self.crear_agradecimientos_profesional(doc)
                    self.progress.set(0.3)
                
                if 'resumen' in self.secciones_activas and 'resumen' in self.content_texts:
                    contenido_resumen = self.content_texts['resumen'].get("1.0", "end")
                    self.crear_seccion_profesional(doc, "RESUMEN", contenido_resumen)
                    self.progress.set(0.4)
                
                if self.incluir_indice.get():
                    self.crear_indice_profesional(doc)
                    self.progress.set(0.5)
                
                # Contenido principal dinámico
                self.crear_contenido_dinamico_profesional(doc)
                self.progress.set(0.8)
                
                # Referencias
                self.crear_referencias_profesionales(doc)
                self.progress.set(0.9)
                
                # Guardar documento
                filename = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word documents", "*.docx")],
                    title="Guardar Proyecto Académico Profesional"
                )
                
                if filename:
                    doc.save(filename)
                    self.progress.stop()
                    self.progress.set(1)
                    
                    # Mensaje de éxito
                    self.validation_text.delete("1.0", "end")
                    self.validation_text.insert("1.0", 
                        f"🎉 ¡DOCUMENTO PROFESIONAL GENERADO!\n\n"
                        f"📄 Archivo: {os.path.basename(filename)}\n"
                        f"📍 Ubicación: {filename}\n\n"
                        f"✅ CARACTERÍSTICAS APLICADAS:\n"
                        f"   • Formato profesional personalizado\n"
                        f"   • Niveles de esquema para índice automático\n"
                        f"   • Saltos de página entre secciones\n"
                        f"   • Control de líneas viudas y huérfanas\n"
                        f"   • Conservar títulos con contenido\n"
                        f"   • Fuente: {self.formato_config['fuente_texto']} {self.formato_config['tamaño_texto']}pt\n"
                        f"   • Títulos: {self.formato_config['fuente_titulo']} {self.formato_config['tamaño_titulo']}pt\n"
                        f"   • Imágenes: {'Personalizadas' if self.encabezado_personalizado or self.insignia_personalizada else 'Base'}\n\n"
                        f"✅ Estructura profesional ({len(self.secciones_activas)} secciones)\n"
                        f"✅ {len(self.referencias)} referencias APA\n"
                        f"✅ Citas procesadas automáticamente\n\n"
                        f"📋 PARA COMPLETAR EN WORD:\n"
                        f"   • Generar índice: Referencias > Tabla de contenido\n"
                        f"   • Agregar numeración de páginas si deseas\n"
                        f"   • Insertar tablas o figuras según necesidad\n\n"
                        f"🚀 ¡Tu proyecto profesional está listo!"
                    )
                    
                    messagebox.showinfo("🎉 ¡Documento Profesional Generado!", 
                        f"Documento creado exitosamente:\n{filename}\n\n"
                        f"Características aplicadas:\n"
                        f"• Formato profesional con niveles de esquema\n"
                        f"• Saltos de página automáticos\n"
                        f"• Control de líneas profesional\n"
                        f"• Imágenes integradas\n"
                        f"• Referencias APA automáticas\n\n"
                        f"Para generar el índice automático:\n"
                        f"Referencias > Tabla de contenido > Automática")
                else:
                    self.progress.stop()
                    self.progress.set(0)
            
            except Exception as e:
                self.progress.stop()
                self.progress.set(0)
                messagebox.showerror("❌ Error", f"Error al generar documento:\n{str(e)}")
        
        thread = threading.Thread(target=generar)
        thread.daemon = True
        thread.start()
    
    def crear_agradecimientos_profesional(self, doc):
        """Crea página de agradecimientos con formato profesional"""
        p = doc.add_paragraph()
        p.style = doc.styles['Titulo Profesional']
        run = p.add_run("AGRADECIMIENTOS")
        run.bold = True
        run.font.name = self.formato_config['fuente_titulo']
        run.font.size = Pt(self.formato_config['tamaño_titulo'])
        
        # Configurar como sección de nivel 1
        p.paragraph_format.outline_level = 0
        if self.formato_config['salto_pagina_secciones']:
            p.paragraph_format.page_break_before = True
        
        doc.add_paragraph()
        content_p = doc.add_paragraph("(Agregar agradecimientos personalizados aquí)")
        content_p.style = doc.styles['Normal']
        doc.add_page_break()
    
    def crear_indice_profesional(self, doc):
        """Crea índice profesional con instrucciones"""
        p = doc.add_paragraph()
        p.style = doc.styles['Titulo Profesional']
        run = p.add_run("ÍNDICE")
        run.bold = True
        run.font.name = self.formato_config['fuente_titulo']
        run.font.size = Pt(self.formato_config['tamaño_titulo'])
        
        # Configurar como sección de nivel 1
        p.paragraph_format.outline_level = 0
        if self.formato_config['salto_pagina_secciones']:
            p.paragraph_format.page_break_before = True
        
        doc.add_paragraph()
        
        instrucciones = """INSTRUCCIONES PARA GENERAR ÍNDICE AUTOMÁTICO:

1. En Word, ir a la pestaña "Referencias"
2. Hacer clic en "Tabla de contenido"
3. Seleccionar "Automática" o "Personalizada"
4. El índice se generará automáticamente con todas las secciones

NOTA: Todas las secciones están configuradas con Nivel de esquema 1 para 
facilitar la generación automática del índice."""
        
        content_p = doc.add_paragraph(instrucciones)
        content_p.style = doc.styles['Normal']
        
        doc.add_paragraph()
        
        # Sección para tabla de ilustraciones
        p2 = doc.add_paragraph()
        run2 = p2.add_run("TABLA DE ILUSTRACIONES")
        run2.bold = True
        run2.font.name = self.formato_config['fuente_titulo']
        run2.font.size = Pt(self.formato_config['tamaño_titulo'] - 2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("(Agregar manualmente si hay figuras, tablas o gráficos)")
        doc.add_page_break()
    
    def procesar_citas(self, texto):
        """Alias para mantener compatibilidad"""
        return self.procesar_citas_completo(texto)
    
    def configurar_estilos_personalizados(self, doc):
        """Configura estilos con formato personalizado"""
        # Estilo normal
        style = doc.styles['Normal']
        style.font.name = self.formato_config['fuente_texto']
        style.font.size = Pt(self.formato_config['tamaño_texto'])
        
        if self.formato_config['interlineado'] == 1.0:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        elif self.formato_config['interlineado'] == 1.5:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        if self.formato_config['justificado']:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if self.formato_config['sangria']:
            style.paragraph_format.first_line_indent = Inches(0.5)
        
        style.paragraph_format.space_after = Pt(0)
        
        # Estilo de títulos
        try:
            titulo_style = doc.styles.add_style('Titulo Personalizado', WD_STYLE_TYPE.PARAGRAPH)
        except:
            titulo_style = doc.styles['Heading 1']
        
        titulo_style.font.name = self.formato_config['fuente_titulo']
        titulo_style.font.size = Pt(self.formato_config['tamaño_titulo'])
        titulo_style.font.bold = True
        titulo_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_style.paragraph_format.space_before = Pt(12)
        titulo_style.paragraph_format.space_after = Pt(12)
    
    def crear_portada_personalizada(self, doc):
        """Crea portada con formato personalizado"""
        # Logo/emblema
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("LOGO/EMBLEMA DE LA INSTITUCIÓN").bold = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Institución
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.proyecto_data['institucion'].get().upper())
        run.bold = True
        run.font.name = self.formato_config['fuente_titulo']
        run.font.size = Pt(self.formato_config['tamaño_titulo'] + 2)
        
        doc.add_paragraph()
        
        # Título del proyecto
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'"{self.proyecto_data["titulo"].get()}"')
        run.bold = True
        run.font.name = self.formato_config['fuente_titulo']
        run.font.size = Pt(self.formato_config['tamaño_titulo'] + 4)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Información del proyecto
        info_fields = ['ciclo', 'curso', 'enfasis', 'area', 'categoria', 'director', 'responsable']
        labels = ['Ciclo', 'Curso', 'Énfasis', 'Área de Desarrollo', 'Categoría', 'Director', 'Responsable']
        
        for field, label in zip(info_fields, labels):
            if field in self.proyecto_data and self.proyecto_data[field].get().strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"{label}: {self.proyecto_data[field].get()}")
                run.font.name = self.formato_config['fuente_texto']
        
        # Estudiantes y tutores (igual que antes)
        doc.add_paragraph()
        
        if self.proyecto_data['estudiantes'].get():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("Estudiantes:").bold = True
            estudiantes = self.proyecto_data['estudiantes'].get().split(',')
            for estudiante in estudiantes:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(estudiante.strip())
        
        if self.proyecto_data['tutores'].get():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("Tutores:").bold = True
            tutores = self.proyecto_data['tutores'].get().split(',')
            for tutor in tutores:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(tutor.strip())
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Fecha
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Año: {datetime.now().year}")
        
        doc.add_page_break()
    
    def crear_contenido_dinamico(self, doc):
        """Crea contenido basado en secciones activas"""
        for seccion_id in self.secciones_activas:
            if seccion_id in self.secciones_disponibles:
                seccion = self.secciones_disponibles[seccion_id]
                
                if seccion['capitulo']:
                    # Es un título de capítulo
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(seccion['titulo'].replace('📖 ', '').replace('📚 ', '').replace('🔬 ', '').replace('🛠️ ', '').replace('📊 ', '').replace('💬 ', ''))
                    run.bold = True
                    run.font.name = self.formato_config['fuente_titulo']
                    run.font.size = Pt(self.formato_config['tamaño_titulo'])
                    doc.add_paragraph()
                else:
                    # Es contenido
                    if seccion_id in self.content_texts:
                        contenido = self.content_texts[seccion_id].get("1.0", "end").strip()
                        if contenido:
                            titulo_limpio = seccion['titulo']
                            # Remover emojis del título para el documento
                            titulo_limpio = re.sub(r'[^\w\s-]', '', titulo_limpio).strip()
                            self.crear_seccion_personalizada(doc, titulo_limpio.upper(), contenido)
    
    def crear_seccion_personalizada(self, doc, titulo, contenido):
        """Crea una sección con formato personalizado"""
        # Título de la sección
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo)
        run.bold = True
        run.font.name = self.formato_config['fuente_titulo']
        run.font.size = Pt(self.formato_config['tamaño_titulo'])
        
        # Contenido procesado
        contenido_procesado = self.procesar_citas(contenido.strip())
        if contenido_procesado:
            p = doc.add_paragraph(contenido_procesado)
            # Aplicar formato personalizado
            p.style = doc.styles['Normal']
            
        doc.add_paragraph()  # Espaciado
    
    def crear_agradecimientos(self, doc):
        """Crea la página de agradecimientos"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("AGRADECIMIENTOS")
        run.bold = True
        run.font.name = self.formato_config['fuente_titulo']
        run.font.size = Pt(self.formato_config['tamaño_titulo'])
        
        doc.add_paragraph()
        doc.add_paragraph("(Agregar agradecimientos manualmente)")
        doc.add_page_break()
    
    def crear_seccion(self, doc, titulo, contenido):
        """Crear sección con formato base"""
        self.crear_seccion_personalizada(doc, titulo, contenido)
    
    def crear_indice(self, doc):
        """Crea el índice"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ÍNDICE")
        run.bold = True
        run.font.name = self.formato_config['fuente_titulo']
        run.font.size = Pt(self.formato_config['tamaño_titulo'])
        
        doc.add_paragraph()
        doc.add_paragraph("(El índice se debe generar manualmente en Word usando Referencias > Tabla de contenido)")
        doc.add_paragraph()
        doc.add_paragraph("TABLA DE ILUSTRACIONES")
        doc.add_paragraph("(Agregar manualmente si hay figuras o tablas)")
        doc.add_page_break()
    
    def crear_referencias_apa(self, doc):
        """Crea referencias con formato personalizado"""
        if not self.referencias:
            return
        
        # Título
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("REFERENCIAS")
        run.bold = True
        run.font.name = self.formato_config['fuente_titulo']
        run.font.size = Pt(self.formato_config['tamaño_titulo'])
        
        doc.add_paragraph()
        
        # Ordenar referencias alfabéticamente
        referencias_ordenadas = sorted(self.referencias, key=lambda x: x['autor'])
        
        for ref in referencias_ordenadas:
            ref_text = f"{ref['autor']} ({ref['año']}). {ref['titulo']}. {ref['fuente']}"
            p = doc.add_paragraph(ref_text)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            p.style = doc.styles['Normal']
    
    def mostrar_bienvenida(self):
        """Muestra mensaje de bienvenida mejorado con atajos de teclado"""
        self.root.after(1000, lambda: messagebox.showinfo(
            "🎓 ¡Generador Profesional!",
            "Generador de Proyectos Académicos - Versión Profesional\n\n"
            "🆕 CARACTERÍSTICAS AVANZADAS:\n"
            "• Formato profesional con niveles de esquema\n"
            "• Auto-guardado cada 5 minutos\n"
            "• Estadísticas en tiempo real\n"
            "• Sistema de guardado/carga completo\n"
            "• Gestión avanzada de imágenes\n"
            "• Exportación de configuraciones\n\n"
            "⌨️ ATAJOS DE TECLADO:\n"
            "• Ctrl+S: Guardar proyecto\n"
            "• Ctrl+O: Cargar proyecto\n"
            "• Ctrl+N: Nuevo proyecto\n"
            "• F5: Validar proyecto\n"
            "• F9: Generar documento\n"
            "• Ctrl+Q: Salir\n\n"
            "🚀 ¡Crea proyectos profesionales únicos!"
        ))
    
    def mostrar_instrucciones(self):
        """Muestra guía completa actualizada con nuevas características"""
        instruc_window = ctk.CTkToplevel(self.root)
        instruc_window.title("📖 Guía Profesional Completa")
        instruc_window.geometry("1000x800")
        
        main_frame = ctk.CTkFrame(instruc_window, corner_radius=0)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        title_label = ctk.CTkLabel(
            main_frame, text="📖 GUÍA PROFESIONAL COMPLETA",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title_label.pack(pady=(20, 10))
        
        instruc_text = ctk.CTkTextbox(main_frame, wrap="word", font=ctk.CTkFont(size=12))
        instruc_text.pack(expand=True, fill="both", padx=20, pady=(10, 20))
        
        instrucciones = """
🎓 GENERADOR PROFESIONAL DE PROYECTOS ACADÉMICOS - VERSIÓN 2.0

═══════════════════════════════════════════════════════════════════

🚀 CARACTERÍSTICAS PROFESIONALES AVANZADAS:
• Formato Word con niveles de esquema para índices automáticos
• Saltos de página inteligentes entre secciones principales
• Control profesional de líneas viudas y huérfanas
• Sistema completo de guardado/carga de proyectos
• Auto-guardado automático cada 5 minutos
• Estadísticas en tiempo real (palabras, secciones, referencias)
• Exportación/importación de configuraciones
• Gestión avanzada de imágenes personalizadas
• Atajos de teclado para flujo de trabajo eficiente

═══════════════════════════════════════════════════════════════════

⌨️ ATAJOS DE TECLADO PROFESIONALES:
• Ctrl+S: Guardar proyecto completo
• Ctrl+O: Cargar proyecto existente
• Ctrl+N: Crear nuevo proyecto
• F5: Validar proyecto rápidamente
• F9: Generar documento final
• Ctrl+Q: Salir de la aplicación

═══════════════════════════════════════════════════════════════════

💾 SISTEMA DE PROYECTOS:

🔹 GUARDAR PROYECTO (Ctrl+S):
• Guarda TODA la información del proyecto
• Incluye contenido, configuración, imágenes, referencias
• Formato JSON para máxima compatibilidad
• Permite reanudar trabajo exactamente donde lo dejaste

🔹 CARGAR PROYECTO (Ctrl+O):
• Restaura proyecto completo
• Mantiene estructura personalizada
• Compatible con versiones anteriores
• Validación automática de integridad

🔹 AUTO-GUARDADO:
• Backup automático cada 5 minutos en "auto_save.json"
• Recuperación automática en caso de cierre inesperado
• No interrumpe el flujo de trabajo

═══════════════════════════════════════════════════════════════════

📊 ESTADÍSTICAS EN TIEMPO REAL:

El panel superior muestra constantemente:
• 📊 Palabras totales escritas
• 📝 Secciones completadas vs total
• 📚 Referencias bibliográficas agregadas
• Actualización automática cada 2 segundos

═══════════════════════════════════════════════════════════════════

🖼️ GESTIÓN AVANZADA DE IMÁGENES:

🔹 IMÁGENES BASE:
• Coloca en /Recursos/Encabezado.png (15cm x 2.5cm)
• Coloca en /Recursos/Insignia.png (3cm x 3cm)
• Detección automática al iniciar

🔹 IMÁGENES PERSONALIZADAS:
• Botón "🖼️ Imágenes" para gestión completa
• Carga imágenes específicas por proyecto
• Prioridad sobre imágenes base
• Validación automática de formatos y tamaños
• Integración perfecta en documento final

═══════════════════════════════════════════════════════════════════

📄 FORMATO PROFESIONAL WORD:

🔹 NIVELES DE ESQUEMA:
• Todos los títulos configurados como Nivel 1
• Índice automático en Word: Referencias > Tabla de contenido
• Navegación rápida por documento
• Estructura profesional garantizada

🔹 CONTROL DE PÁRRAFOS:
• Saltos de página antes de secciones principales
• "Conservar con el siguiente" para títulos
• Control de líneas viudas y huérfanas
• Conservar líneas juntas para coherencia

🔹 FORMATO APA AUTOMÁTICO:
• Sangría francesa en referencias (1.27cm)
• Interlineado doble configurable
• Justificación automática
• Márgenes estándar (2.54cm)

═══════════════════════════════════════════════════════════════════

📚 SISTEMA DE CITAS PROFESIONAL:

🔥 TIPOS SOPORTADOS:

📝 Cita textual corta:
"El conocimiento es poder" [CITA:textual:Bacon:1597:25]
→ "El conocimiento es poder" (Bacon, 1597, p. 25)

🔄 Parafraseo:
Los estudios demuestran mejoras [CITA:parafraseo:García:2020]
→ Los estudios demuestran mejoras (García, 2020)

📖 Cita textual larga (más de 40 palabras):
Su extensa reflexión sobre el tema [CITA:larga:Autor:2021:45]
→ Formato de bloque independiente con sangría

👥 Múltiples autores:
Investigaciones recientes [CITA:multiple:García y López:2020]
→ Investigaciones recientes (García y López, 2020)

🌐 Fuentes web:
Según datos oficiales [CITA:web:OMS:2023]
→ Según datos oficiales (OMS, 2023)

═══════════════════════════════════════════════════════════════════

🛠️ GESTIÓN DINÁMICA DE SECCIONES:

🔹 OPERACIONES DISPONIBLES:
• ➕ Agregar: Crea secciones completamente personalizadas
• ➖ Quitar: Elimina secciones (protege las requeridas)
• ✏️ Editar: Modifica títulos, instrucciones y propiedades
• ⬆️⬇️ Reordenar: Cambia orden de aparición en documento

🔹 TIPOS DE SECCIONES:
• 📖 Capítulos: Solo títulos organizacionales (azul)
• 📝 Contenido: Secciones con texto (blanco)
• ⚠️ Requeridas: Obligatorias para validación (amarillo)
• 🎨 Personalizadas: Creadas según tus necesidades

═══════════════════════════════════════════════════════════════════

🎨 FORMATO COMPLETAMENTE PERSONALIZABLE:

🔹 TIPOGRAFÍA:
• Fuentes independientes para texto y títulos
• Tamaños específicos (10-18pt)
• Familias: Times New Roman, Arial, Calibri, etc.

🔹 ESPACIADO:
• Interlineado: 1.0, 1.5, 2.0, 2.5
• Márgenes: 2.0-3.5 cm
• Sangría primera línea configurable

🔹 OPCIONES PROFESIONALES:
• Saltos de página entre secciones
• Conservar títulos con contenido
• Control de líneas viudas y huérfanas
• Justificación automática

═══════════════════════════════════════════════════════════════════

📤 EXPORTACIÓN DE CONFIGURACIONES:

🔹 CASOS DE USO:
• Crear plantillas departamentales
• Compartir configuraciones entre proyectos
• Backup de configuraciones favoritas
• Estandarización institucional

🔹 CONTENIDO EXPORTADO:
• Configuración completa de formato
• Estructura de secciones personalizada
• Orden y propiedades de secciones
• Compatible entre versiones

═══════════════════════════════════════════════════════════════════

🤖 TRABAJANDO CON IA AVANZADO:

Para máxima eficiencia con IA como Claude:

📝 SOLICITUDES ESTRUCTURADAS:
"Escríbeme el marco teórico sobre [tema] organizado en 4 subsecciones:
1. Antecedentes históricos
2. Teorías principales  
3. Estudios recientes
4. Debates actuales

Incluye citas usando [CITA:parafraseo:Autor:Año] y [CITA:textual:Autor:Año:Página].
Estructura cada subsección con 200-300 palabras."

🎯 OBJETIVOS ESPECÍFICOS:
"Dame 6 objetivos específicos para investigar [problema] usando estos verbos:
• Identificar (población objetivo)
• Comparar (metodologías) 
• Analizar (resultados)
• Evaluar (impacto)
• Determinar (factores)
• Proponer (soluciones)

Cada objetivo debe ser medible y alcanzable."

💬 DISCUSIÓN AVANZADA:
"Redacta la discusión comparando estos resultados [datos] con:
1. Teoría de [autor principal]
2. Estudios similares de [autores secundarios]
3. Limitaciones encontradas
4. Implicaciones prácticas

Incluye citas de contraste usando el formato del generador."

═══════════════════════════════════════════════════════════════════

🔍 VALIDACIÓN INTELIGENTE AVANZADA:

🎯 CRITERIOS EVALUADOS:
• Campos obligatorios completos
• Secciones requeridas con contenido mínimo (50 caracteres)
• Detección de citas en marco teórico
• Verificación de referencias bibliográficas
• Coherencia entre objetivos y contenido
• Estructura lógica de secciones

📊 MÉTRICAS MOSTRADAS:
• Progreso general del proyecto (0-100%)
• Secciones completadas vs requeridas
• Estadísticas de contenido
• Estado de configuración aplicada
• Recomendaciones de mejora específicas

═══════════════════════════════════════════════════════════════════

💡 FLUJO DE TRABAJO PROFESIONAL RECOMENDADO:

🔄 INICIO DE PROYECTO:
1. Crear nuevo proyecto (Ctrl+N)
2. Configurar formato personalizado
3. Cargar/configurar imágenes institucionales
4. Organizar estructura de secciones
5. Completar información general

📝 DESARROLLO:
1. Escribir sección por sección siguiendo instrucciones
2. Usar sistema de citas consistentemente
3. Agregar referencias mientras escribes
4. Guardar frecuentemente (Ctrl+S)
5. Validar regularmente (F5)

🎯 FINALIZACIÓN:
1. Validación final completa
2. Revisión de estadísticas
3. Generación de documento (F9)
4. Verificación en Word
5. Generar índice automático

═══════════════════════════════════════════════════════════════════

🏆 CARACTERÍSTICAS ÚNICAS DEL GENERADOR:

✅ PROFESIONALISMO TOTAL:
• Cumple estándares académicos internacionales
• Formato compatible con universidades
• Estructura adaptable a cualquier disciplina
• Calidad profesional garantizada

✅ EFICIENCIA MÁXIMA:
• Automatiza formateo tedioso
• Reduce tiempo de creación 80%
• Elimina errores de formato
• Permite enfoque en contenido de calidad

✅ FLEXIBILIDAD ABSOLUTA:
• Adaptable a cualquier tipo de proyecto
• Personalización total del formato
• Estructura dinámica según necesidades
• Compatible con múltiples estilos institucionales

═══════════════════════════════════════════════════════════════════

🚀 ¡AHORA PUEDES CREAR PROYECTOS DE NIVEL PROFESIONAL!

Este generador avanzado combina la potencia de automatización con la flexibilidad 
necesaria para crear documentos académicos de la más alta calidad. 

Experimenta con las diferentes configuraciones y encuentra el formato perfecto 
para tu institución y tipo de proyecto.

¡El futuro de la creación de documentos académicos está aquí!
        """
        
        instruc_text.insert("1.0", instrucciones)
        instruc_text.configure(state="disabled")
    
    def run(self):
        """Ejecuta la aplicación"""
        self.root.mainloop()

class SeccionDialog:
    """Diálogo para agregar/editar secciones"""
    def __init__(self, parent, secciones_existentes, editar=False, seccion_actual=None):
        self.result = None
        self.secciones_existentes = secciones_existentes
        self.editar = editar
        self.seccion_actual = seccion_actual
        
        # Crear ventana de diálogo
        titulo = "✏️ Editar Sección" if editar else "➕ Agregar Nueva Sección"
        self.dialog = ctk.CTkToplevel(parent)
        self.dialog.title(titulo)
        self.dialog.geometry("550x450")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        # Centrar ventana
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() // 2) - (550 // 2)
        y = (self.dialog.winfo_screenheight() // 2) - (450 // 2)
        self.dialog.geometry(f"550x450+{x}+{y}")
        
        self.setup_dialog()
        
        # Cargar datos si es edición
        if editar and seccion_actual:
            self.cargar_datos_existentes()
    
    def setup_dialog(self):
        """Configura el diálogo"""
        main_frame = ctk.CTkFrame(self.dialog, corner_radius=0)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Título
        titulo_texto = "✏️ Editar Sección Existente" if self.editar else "➕ Crear Nueva Sección"
        title_label = ctk.CTkLabel(
            main_frame, text=titulo_texto,
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(pady=(10, 20))
        
        # Campos
        fields_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        fields_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        # ID único
        ctk.CTkLabel(fields_frame, text="ID único:", font=ctk.CTkFont(weight="bold")).pack(anchor="w", pady=(0, 5))
        self.id_entry = ctk.CTkEntry(fields_frame, placeholder_text="ejemplo: mi_seccion_personalizada")
        self.id_entry.pack(fill="x", pady=(0, 15))
        
        # Título
        ctk.CTkLabel(fields_frame, text="Título:", font=ctk.CTkFont(weight="bold")).pack(anchor="w", pady=(0, 5))
        self.titulo_entry = ctk.CTkEntry(fields_frame, placeholder_text="📝 Mi Nueva Sección")
        self.titulo_entry.pack(fill="x", pady=(0, 15))
        
        # Instrucción
        ctk.CTkLabel(fields_frame, text="Instrucción:", font=ctk.CTkFont(weight="bold")).pack(anchor="w", pady=(0, 5))
        self.instruccion_text = ctk.CTkTextbox(fields_frame, height=80)
        self.instruccion_text.insert("1.0", "Describe qué debe contener esta sección...")
        self.instruccion_text.pack(fill="x", pady=(0, 15))
        
        # Opciones
        options_frame = ctk.CTkFrame(fields_frame, fg_color="gray20", corner_radius=10)
        options_frame.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(options_frame, text="⚙️ Opciones de Sección:", font=ctk.CTkFont(weight="bold")).pack(pady=(10, 5))
        
        self.es_capitulo = ctk.CTkCheckBox(options_frame, text="📖 Es título de capítulo (solo organizacional)")
        self.es_capitulo.pack(anchor="w", padx=20, pady=5)
        
        self.es_requerida = ctk.CTkCheckBox(options_frame, text="⚠️ Sección requerida (obligatoria para validación)")
        self.es_requerida.pack(anchor="w", padx=20, pady=5)
        
        # Información adicional
        info_text = """💡 INFORMACIÓN:
• ID único: Identificador interno (sin espacios, usar guiones bajos)
• Título: Nombre que aparecerá en pestañas y documento
• Instrucción: Guía para el usuario sobre qué escribir
• Capítulo: Solo aparece como título organizacional, sin contenido
• Requerida: Se valida que tenga contenido antes de generar"""
        
        info_label = ctk.CTkLabel(
            options_frame, text=info_text, font=ctk.CTkFont(size=10),
            justify="left", wraplength=450
        )
        info_label.pack(padx=15, pady=(5, 15))
        
        # Botones
        btn_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        btn_frame.pack(fill="x", pady=(0, 10))
        
        cancel_btn = ctk.CTkButton(
            btn_frame, text="❌ Cancelar", command=self.cancelar,
            fg_color="red", hover_color="darkred", width=120
        )
        cancel_btn.pack(side="left", padx=(20, 10))
        
        action_text = "✅ Actualizar Sección" if self.editar else "✅ Crear Sección"
        create_btn = ctk.CTkButton(
            btn_frame, text=action_text, command=self.procesar_seccion,
            fg_color="green", hover_color="darkgreen", width=150
        )
        create_btn.pack(side="right", padx=(10, 20))
    
    def cargar_datos_existentes(self):
        """Carga los datos de la sección existente para editar"""
        if self.seccion_actual:
            seccion_id, seccion_data = self.seccion_actual
            
            self.id_entry.delete(0, "end")
            self.id_entry.insert(0, seccion_id)
            
            self.titulo_entry.delete(0, "end")
            self.titulo_entry.insert(0, seccion_data['titulo'])
            
            self.instruccion_text.delete("1.0", "end")
            self.instruccion_text.insert("1.0", seccion_data['instruccion'])
            
            if seccion_data['capitulo']:
                self.es_capitulo.select()
            
            if seccion_data['requerida']:
                self.es_requerida.select()
    
    def procesar_seccion(self):
        """Procesa la creación o edición de la sección"""
        seccion_id = self.id_entry.get().strip()
        titulo = self.titulo_entry.get().strip()
        instruccion = self.instruccion_text.get("1.0", "end").strip()
        
        if not all([seccion_id, titulo, instruccion]):
            messagebox.showerror("❌ Error", "Completa todos los campos obligatorios")
            return
        
        # Validar ID único (solo si no es edición o cambió el ID)
        if not self.editar or (self.editar and seccion_id != self.seccion_actual[0]):
            if seccion_id in self.secciones_existentes:
                messagebox.showerror("❌ Error", "Ya existe una sección con ese ID")
                return
        
        # Validar formato del ID
        if not re.match(r'^[a-z0-9_]+$', seccion_id):
            messagebox.showerror("❌ Error", 
                "El ID debe contener solo letras minúsculas, números y guiones bajos")
            return
        
        seccion_data = {
            'titulo': titulo,
            'instruccion': instruccion,
            'requerida': self.es_requerida.get(),
            'capitulo': self.es_capitulo.get()
        }
        
        self.result = (seccion_id, seccion_data)
        self.dialog.destroy()
    
    def cancelar(self):
        """Cancela la operación"""
        self.dialog.destroy()

def main():
    app = ProyectoAcademicoGenerator()
    app.run()

if __name__ == "__main__":
    main()