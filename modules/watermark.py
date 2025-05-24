"""
Sistema Avanzado de Marcas de Agua para Documentos - Versión Actualizada
Compatible con diferentes versiones de python-docx
Incluye configuraciones específicas de posición y tamaño
"""

import os
from PIL import Image, ImageEnhance
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64

class WatermarkManager:
    def __init__(self):
        self.default_opacity = 0.3
        self.default_position = 'header'
        self.cache = {}
        
        # Configuración del encabezado (páginas 2 en adelante)
        self.header_config = {
            'width': Cm(20.96),
            'height': Cm(27.68),
            'h_align': 'center',      # Centrado horizontalmente
            'v_position': Cm(-1.5),   # -1.5 cm desde el párrafo
            'behind_text': True       # Detrás del texto
        }
        
        # Configuración de la insignia (solo primera página)
        self.logo_config = {
            'width': Cm(6.11),
            'height': Cm(7.45),
            'align': 'center'  # Centrada
        }
        
    def process_image_for_watermark(self, image_path, opacity=None, width_inches=None):
        """Procesa una imagen para usarla como marca de agua"""
        if opacity is None:
            opacity = self.default_opacity
            
        cache_key = f"{image_path}_{opacity}_{width_inches}"
        if cache_key in self.cache:
            return self.cache[cache_key]
        
        try:
            # Abrir imagen
            img = Image.open(image_path)
            
            # Convertir a RGBA si es necesario
            if img.mode != 'RGBA':
                img = img.convert('RGBA')
            
            # Redimensionar si se especifica ancho
            if width_inches:
                # Convertir pulgadas a píxeles (asumiendo 96 DPI)
                width_px = int(width_inches * 96)
                ratio = width_px / img.width
                height_px = int(img.height * ratio)
                img = img.resize((width_px, height_px), Image.Resampling.LANCZOS)
            
            # Aplicar transparencia
            if img.mode == 'RGBA':
                # Procesar canal alpha
                alpha = img.split()[-1]
                alpha = ImageEnhance.Brightness(alpha).enhance(opacity)
                img.putalpha(alpha)
            else:
                # Si no tiene alpha, crear uno
                img = img.convert('RGBA')
                data = img.getdata()
                newData = []
                for item in data:
                    # Cambiar todos los píxeles blancos (o casi blancos) a transparentes
                    if len(item) == 4:
                        newData.append((item[0], item[1], item[2], int(item[3] * opacity)))
                    else:
                        newData.append((item[0], item[1], item[2], int(255 * opacity)))
                img.putdata(newData)
            
            # Guardar en memoria
            buffer = io.BytesIO()
            img.save(buffer, format='PNG')
            buffer.seek(0)
            
            self.cache[cache_key] = buffer.getvalue()
            return self.cache[cache_key]
            
        except Exception as e:
            print(f"Error procesando imagen para marca de agua: {e}")
            return None
    
    def configurar_imagen_detras_texto(self, picture, config=None):
        """Configura una imagen para que aparezca detrás del texto con posición específica"""
        if config is None:
            config = self.header_config
            
        inline = picture._inline
        
        # Crear elemento anchor
        anchor = OxmlElement('wp:anchor')
        
        # Atributos básicos
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '0')
        anchor.set('distR', '0')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '0')
        anchor.set('behindDoc', '1')  # Detrás del texto
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')
        
        # Posición simple (requerida)
        simplePos = OxmlElement('wp:simplePos')
        simplePos.set('x', '0')
        simplePos.set('y', '0')
        anchor.append(simplePos)
        
        # Posición horizontal - Centrada con relación a página
        positionH = OxmlElement('wp:positionH')
        positionH.set('relativeFrom', 'page')
        alignH = OxmlElement('wp:align')
        alignH.text = 'center'
        positionH.append(alignH)
        anchor.append(positionH)
        
        # Posición vertical - Absoluta desde párrafo
        positionV = OxmlElement('wp:positionV')
        positionV.set('relativeFrom', 'paragraph')
        posOffsetV = OxmlElement('wp:posOffset')
        posOffsetV.text = str(int(config.get('v_position', Cm(-1.5))))
        positionV.append(posOffsetV)
        anchor.append(positionV)
        
        # Tamaño
        extent = OxmlElement('wp:extent')
        extent.set('cx', str(int(config.get('width', Cm(20.96)))))
        extent.set('cy', str(int(config.get('height', Cm(27.68)))))
        anchor.append(extent)
        
        # Efecto visual (ninguno = detrás del texto)
        wrapNone = OxmlElement('wp:wrapNone')
        anchor.append(wrapNone)
        
        # Copiar elementos del inline al anchor
        for element in inline:
            if element.tag.endswith(('docPr', 'cNvGraphicFramePr', 'graphic')):
                anchor.append(element)
                
        # Reemplazar inline con anchor
        inline.getparent().replace(inline, anchor)
    
    def add_watermark_to_section(self, section, image_path, opacity=0.3, stretch=True):
        """Agrega marca de agua a una sección del documento con configuración mejorada"""
        try:
            # Configurar sección
            section.header_distance = Cm(1.25)  # Margen del encabezado
            section.footer_distance = Cm(1.25)  # Margen del pie de página
            
            # Obtener header
            header = section.header
            
            # Limpiar header existente si es necesario
            if not header.paragraphs:
                header_para = header.add_paragraph()
            else:
                header_para = header.paragraphs[0]
            
            # Configurar alineación
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Agregar la imagen
            run = header_para.add_run()
            
            # Usar configuración específica
            header_pic = run.add_picture(image_path, width=self.header_config['width'])
            
            # Configurar posición detrás del texto
            try:
                self.configurar_imagen_detras_texto(header_pic, self.header_config)
                print("✅ Encabezado configurado con posicionamiento avanzado")
                return True
            except Exception as e:
                print(f"⚠️ Error en posicionamiento avanzado: {e}")
                # Intentar método alternativo
                return self._add_watermark_alternative(header_para, image_path, opacity, stretch)
                    
        except Exception as e:
            print(f"Error agregando marca de agua: {e}")
            return False
    
    def add_logo_to_first_page(self, doc, logo_path):
        """Agrega insignia/logo centrado en la primera página"""
        if not logo_path or not os.path.exists(logo_path):
            return False
            
        try:
            # Agregar párrafo para el logo
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Agregar imagen con configuración específica
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path, height=self.logo_config['height'])
            
            # Espacio después del logo
            doc.add_paragraph()
            
            print("✅ Insignia agregada a primera página")
            return True
            
        except Exception as e:
            print(f"Error agregando insignia: {e}")
            return False
    
    def _add_watermark_alternative(self, paragraph, image_path, opacity, stretch):
        """Método alternativo para agregar marca de agua usando XML directo"""
        try:
            # Procesar imagen primero
            if stretch:
                # Convertir cm a pulgadas para procesamiento
                width_inches = self.header_config['width'] / Cm(1) / 2.54
                processed_image = self.process_image_for_watermark(image_path, opacity, width_inches)
            else:
                processed_image = self.process_image_for_watermark(image_path, opacity, 5)
            
            if not processed_image:
                return False
            
            # Convertir a base64
            image_base64 = base64.b64encode(processed_image).decode('utf-8')
            
            # Crear elemento run
            run = paragraph.add_run()
            r = run._r
            
            # Crear estructura pict
            pict = OxmlElement('w:pict')
            
            # Crear shape con configuración específica
            shape = OxmlElement('v:shape')
            shape.set('id', '_x0000_i1025')
            shape.set('type', '#_x0000_t75')
            
            # Convertir dimensiones a puntos para el estilo
            width_pt = int(self.header_config['width'] / Cm(1) * 28.35)
            height_pt = int(self.header_config['height'] / Cm(1) * 28.35)
            
            shape.set('style', f'width:{width_pt}pt;height:{height_pt}pt;position:absolute;z-index:-251658752')
            
            # Crear imagedata
            imagedata = OxmlElement('v:imagedata')
            imagedata.set(qn('r:id'), 'rId1')
            imagedata.set('o:title', 'Watermark')
            
            # Agregar imagedata a shape
            shape.append(imagedata)
            
            # Agregar shape a pict
            pict.append(shape)
            
            # Agregar pict a run
            r.append(pict)
            
            # Intentar agregar la relación de imagen
            try:
                # Este es un método simplificado, puede necesitar ajustes
                document = paragraph.part
                image_part = document.new_image_part(image_path)
                imagedata.set(qn('r:id'), document.relate_to(image_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'))
            except:
                pass
            
            return True
            
        except Exception as e:
            print(f"Error en método alternativo: {e}")
            return False
    
    def add_simple_header_image(self, section, image_path, width_inches=None):
        """Método simple para agregar imagen al encabezado con dimensiones correctas"""
        try:
            header = section.header
            
            # Configurar márgenes de sección
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.25)
            
            # Asegurar que hay un párrafo
            if not header.paragraphs:
                p = header.add_paragraph()
            else:
                p = header.paragraphs[0]
            
            # Centrar el párrafo
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Agregar la imagen con configuración específica
            run = p.add_run()
            
            # Usar ancho de configuración si no se especifica
            if width_inches is None:
                picture = run.add_picture(image_path, width=self.header_config['width'])
            else:
                picture = run.add_picture(image_path, width=Inches(width_inches))
            
            return True
            
        except Exception as e:
            print(f"Error agregando imagen simple al header: {e}")
            return False
    
    def configure_document_headers(self, doc, header_image_path, logo_image_path=None):
        """Configura encabezados del documento completo según especificaciones"""
        try:
            # Configurar primera sección
            section = doc.sections[0]
            
            # IMPORTANTE: Primera página diferente
            section.different_first_page_header_footer = True
            
            # Configurar márgenes
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.25)
            
            # Agregar encabezado a páginas 2+
            if header_image_path and os.path.exists(header_image_path):
                self.add_watermark_to_section(section, header_image_path)
            
            print("✅ Configuración de encabezados completada")
            return True
            
        except Exception as e:
            print(f"Error configurando encabezados del documento: {e}")
            return False