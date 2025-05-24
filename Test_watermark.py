#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script final para gestión de encabezados e insignias en documentos Word
Con posicionamiento y configuración corregidos
"""

import os
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image as PILImage

class DocumentoConEncabezados:
    """Gestor para documentos con encabezados e insignias"""
    
    def __init__(self):
        # Configuración del encabezado (páginas 2 en adelante)
        self.header_config = {
            'width': Cm(20.96),
            'height': Cm(27.68),
            'h_align': 'center',      # Centrado horizontalmente
            'v_position': Cm(-1.5),   # -1.5 cm desde el párrafo
            'behind_text': True       # Detrás del texto
        }
        
        # Configuración de la insignia (solo primera página)
        self.logo_config = {
            'width': Cm(6.11),
            'height': Cm(7.45),
            'align': 'center'  # Centrada
        }
        
        # Rutas de imágenes
        self.header_image = None
        self.logo_image = None
        
    def buscar_imagenes(self):
        """Busca las imágenes necesarias"""
        # Buscar encabezado
        posibles_encabezados = [
            'resources/images/Encabezado.png',
            'resources/images/Encabezado.jpg',
            'images/Encabezado.png',
            'Encabezado.png'
        ]
        
        for ruta in posibles_encabezados:
            if os.path.exists(ruta):
                self.header_image = ruta
                print(f"✅ Encabezado encontrado: {ruta}")
                break
                
        # Buscar insignia
        posibles_insignias = [
            'resources/images/Insignia.png',
            'resources/images/Insignia.jpg',
            'images/Insignia.png',
            'Insignia.png',
            'logo.png'
        ]
        
        for ruta in posibles_insignias:
            if os.path.exists(ruta):
                self.logo_image = ruta
                print(f"✅ Insignia encontrada: {ruta}")
                break
                
        if not self.header_image:
            print("❌ No se encontró imagen de encabezado")
        if not self.logo_image:
            print("⚠️ No se encontró imagen de insignia")
            
    def configurar_imagen_detras_texto(self, picture, config):
        """Configura una imagen para que aparezca detrás del texto con posición específica"""
        inline = picture._inline
        
        # Crear elemento anchor
        anchor = OxmlElement('wp:anchor')
        
        # Atributos básicos
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '0')
        anchor.set('distR', '0')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '0')
        anchor.set('behindDoc', '1')  # Detrás del texto
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')
        
        # Posición simple (requerida)
        simplePos = OxmlElement('wp:simplePos')
        simplePos.set('x', '0')
        simplePos.set('y', '0')
        anchor.append(simplePos)
        
        # Posición horizontal - Centrada con relación a página
        positionH = OxmlElement('wp:positionH')
        positionH.set('relativeFrom', 'page')
        alignH = OxmlElement('wp:align')
        alignH.text = 'center'
        positionH.append(alignH)
        anchor.append(positionH)
        
        # Posición vertical - Absoluta desde párrafo
        positionV = OxmlElement('wp:positionV')
        positionV.set('relativeFrom', 'paragraph')
        posOffsetV = OxmlElement('wp:posOffset')
        posOffsetV.text = str(int(config.get('v_position', Cm(-1.5))))
        positionV.append(posOffsetV)
        anchor.append(positionV)
        
        # Tamaño
        extent = OxmlElement('wp:extent')
        extent.set('cx', str(int(config.get('width', Cm(20.96)))))
        extent.set('cy', str(int(config.get('height', Cm(27.68)))))
        anchor.append(extent)
        
        # Efecto visual (ninguno = detrás del texto)
        wrapNone = OxmlElement('wp:wrapNone')
        anchor.append(wrapNone)
        
        # Copiar elementos del inline al anchor
        for element in inline:
            if element.tag.endswith(('docPr', 'cNvGraphicFramePr', 'graphic')):
                anchor.append(element)
                
        # Reemplazar inline con anchor
        inline.getparent().replace(inline, anchor)
        
    def crear_documento_nuevo(self, nombre_archivo="documento_con_encabezados.docx"):
        """Crea un documento nuevo con la configuración correcta"""
        
        # Buscar imágenes
        self.buscar_imagenes()
        
        if not self.header_image:
            print("❌ No se puede continuar sin imagen de encabezado")
            return False
            
        # Crear documento
        doc = Document()
        
        # Configurar estilos
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # Configurar sección
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        section.header_distance = Cm(1.25)  # Margen del encabezado
        section.footer_distance = Cm(1.25)  # Margen del pie de página
        
        # IMPORTANTE: Primera página diferente
        section.different_first_page_header_footer = True
        
        # ===== PRIMERA PÁGINA =====
        print("📄 Configurando primera página...")
        
        # Si hay insignia, agregarla centrada al principio
        if self.logo_image:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run()
            logo_run.add_picture(self.logo_image, height=Cm(7.45))
            # Espacio después de la insignia
            doc.add_paragraph()
        
        # Título principal
        title = doc.add_heading('COLEGIO PRIVADO DIVINA ESPERANZA', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo
        doc.add_paragraph()  # Espacio
        subtitle = doc.add_heading('"TITULO DEL TRABAJO"', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espaciado
        for _ in range(8):
            doc.add_paragraph()
            
        # Información
        info1 = doc.add_paragraph('Ciclo: Tercer año')
        info2 = doc.add_paragraph('Autor: [Nombre del Estudiante]')
        info3 = doc.add_paragraph('Fecha: [Fecha de Entrega]')
        
        # ===== CONFIGURAR ENCABEZADO PARA PÁGINAS 2+ =====
        print("📄 Configurando encabezado para páginas siguientes...")
        
        # Encabezado principal (NO primera página)
        header = section.header
        
        # Agregar imagen del encabezado
        header_para = header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = header_para.add_run()
        header_pic = run.add_picture(self.header_image, width=self.header_config['width'])
        
        # Configurar posición
        try:
            self.configurar_imagen_detras_texto(header_pic, self.header_config)
            print("✅ Encabezado configurado con posicionamiento avanzado")
        except Exception as e:
            print(f"⚠️ Error en posicionamiento avanzado: {e}")
            
        # ===== PÁGINAS ADICIONALES =====
        # Página 2
        doc.add_page_break()
        doc.add_heading('AGRADECIMIENTO', 1)
        agradecimiento = doc.add_paragraph(
            'Lorem ipsum dolor sit amet, consectetur adipiscing elit. '
            'Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. '
            'Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris '
            'nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in '
            'reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur.'
        )
        agradecimiento.paragraph_format.first_line_indent = Cm(1.25)
        agradecimiento.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Más contenido
        for _ in range(3):
            p = doc.add_paragraph(
                'Lorem ipsum dolor sit amet, consectetur adipiscing elit. ' * 5
            )
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        # Página 3
        doc.add_page_break()
        doc.add_heading('ÍNDICE', 1)
        doc.add_paragraph('1. Introducción ........................... 4')
        doc.add_paragraph('2. Marco Teórico ......................... 10')
        doc.add_paragraph('3. Metodología ........................... 25')
        doc.add_paragraph('4. Resultados ............................ 40')
        doc.add_paragraph('5. Conclusiones .......................... 55')
        doc.add_paragraph('6. Referencias ........................... 60')
        
        # Página 4
        doc.add_page_break()
        doc.add_heading('CAPÍTULO 1: INTRODUCCIÓN', 1)
        intro = doc.add_paragraph(
            'Este es el inicio del contenido principal del documento. '
            'En esta página y las siguientes debe aparecer el encabezado '
            'con la imagen del colegio detrás del texto. '
        )
        intro.paragraph_format.first_line_indent = Cm(1.25)
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Más contenido
        for _ in range(5):
            p = doc.add_paragraph(
                'Lorem ipsum dolor sit amet, consectetur adipiscing elit. ' * 8
            )
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        # Guardar documento
        doc.save(nombre_archivo)
        print(f"\n✅ Documento guardado como: {nombre_archivo}")
        print("\n📋 Estructura del documento:")
        print("- Primera página: Insignia centrada arriba + contenido (SIN encabezado)")
        print("- Páginas 2+: Encabezado detrás del texto (SIN insignia)")
        
        return True
        
    def aplicar_a_documento_existente(self, archivo_entrada, archivo_salida=None):
        """Aplica encabezados a un documento existente"""
        
        if not os.path.exists(archivo_entrada):
            print(f"❌ No se encontró el archivo: {archivo_entrada}")
            return False
            
        # Buscar imágenes
        self.buscar_imagenes()
        
        if not self.header_image:
            print("❌ No se puede continuar sin imagen de encabezado")
            return False
            
        # Abrir documento
        doc = Document(archivo_entrada)
        
        # Configurar primera sección
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        
        # Limpiar encabezado principal existente
        header = section.header
        for para in header.paragraphs:
            p = para._element
            p.getparent().remove(p)
            
        # Agregar nuevo encabezado
        header_para = header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = header_para.add_run()
        header_pic = run.add_picture(self.header_image, width=self.header_config['width'])
        
        # Configurar posición
        try:
            self.configurar_imagen_detras_texto(header_pic, self.header_config)
        except Exception as e:
            print(f"⚠️ Error en posicionamiento: {e}")
            
        # Guardar
        if archivo_salida is None:
            archivo_salida = archivo_entrada.replace('.docx', '_con_encabezados.docx')
            
        doc.save(archivo_salida)
        print(f"✅ Documento guardado como: {archivo_salida}")
        
        return True

def main():
    """Función principal"""
    print("=" * 70)
    print("GESTOR DE ENCABEZADOS E INSIGNIAS - VERSIÓN FINAL")
    print("=" * 70)
    
    gestor = DocumentoConEncabezados()
    
    print("\n📋 Configuración:")
    print("- Encabezado: 20.96 x 27.68 cm, centrado, -1.5 cm vertical")
    print("- Insignia: 6.11 x 7.45 cm, centrada en primera página")
    print("- Primera página: SIN encabezado, CON insignia")
    print("- Páginas 2+: CON encabezado, SIN insignia")
    
    print("\n🔧 Creando documento nuevo...")
    gestor.crear_documento_nuevo("documento_final.docx")
    
    print("\n✅ Proceso completado")
    print("\nVerifica que:")
    print("1. La insignia aparece centrada al inicio de la primera página")
    print("2. El encabezado aparece detrás del texto desde la página 2")
    print("3. Los márgenes y posiciones son correctos")

if __name__ == "__main__":
    main()
    input("\nPresiona Enter para salir...")