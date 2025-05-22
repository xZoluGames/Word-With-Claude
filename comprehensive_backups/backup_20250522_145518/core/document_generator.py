"""
Generador de documentos Word - Crea documentos profesionales con formato avanzado
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import threading
import os
from datetime import datetime
from tkinter import filedialog, messagebox
import re

class DocumentGenerator:
    def __init__(self):
        self.formato_config = {
            'fuente_texto': 'Times New Roman',
            'tamaño_texto': 12,
            'fuente_titulo': 'Times New Roman', 
            'tamaño_titulo': 14,
            'interlineado': 2.0,
            'margen': 2.54,
            'justificado': True,
            'sangria': True
        }
    
    def generar_documento_async(self, app_instance):
        """Genera el documento profesional en un hilo separado"""
        def generar():
            try:
                app_instance.progress.set(0)
                app_instance.progress.start()
                
                # Crear documento
                doc = Document()
                
                # Configurar estilos profesionales
                self.configurar_estilos_profesionales(doc, app_instance)
                app_instance.progress.set(0.1)
                
                # Generar contenido
                if app_instance.incluir_portada.get():
                    self.crear_portada_profesional(doc, app_instance)
                    app_instance.progress.set(0.2)
                
                if app_instance.incluir_agradecimientos.get():
                    self.crear_agradecimientos_profesional(doc, app_instance)
                    app_instance.progress.set(0.3)
                
                if 'resumen' in app_instance.secciones_activas and 'resumen' in app_instance.content_texts:
                    contenido_resumen = app_instance.content_texts['resumen'].get("1.0", "end")
                    self.crear_seccion_profesional(doc, "RESUMEN", contenido_resumen, app_instance)
                    app_instance.progress.set(0.4)
                
                if app_instance.incluir_indice.get():
                    self.crear_indice_profesional(doc, app_instance)
                    app_instance.progress.set(0.5)
                
                # Contenido principal dinámico
                self.crear_contenido_dinamico_profesional(doc, app_instance)
                app_instance.progress.set(0.8)
                
                # Referencias
                self.crear_referencias_profesionales(doc, app_instance)
                app_instance.progress.set(0.9)
                
                # Guardar documento
                filename = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word documents", "*.docx")],
                    title="Guardar Proyecto Académico Profesional"
                )
                
                if filename:
                    doc.save(filename)
                    app_instance.progress.stop()
                    app_instance.progress.set(1)
                    
                    self.mostrar_mensaje_exito(filename, app_instance)
                else:
                    app_instance.progress.stop()
                    app_instance.progress.set(0)
            
            except Exception as e:
                app_instance.progress.stop()
                app_instance.progress.set(0)
                messagebox.showerror("❌ Error", f"Error al generar documento:\n{str(e)}")
        
        thread = threading.Thread(target=generar)
        thread.daemon = True
        thread.start()
    
    def configurar_estilos_profesionales(self, doc, app_instance):
        """Configura estilos profesionales del documento"""
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(app_instance.formato_config['margen'] / 2.54)
            section.bottom_margin = Inches(app_instance.formato_config['margen'] / 2.54)
            section.left_margin = Inches(app_instance.formato_config['margen'] / 2.54)
            section.right_margin = Inches(app_instance.formato_config['margen'] / 2.54)
        
        # Estilo normal
        style = doc.styles['Normal']
        style.font.name = app_instance.formato_config['fuente_texto']
        style.font.size = Pt(app_instance.formato_config['tamaño_texto'])
        
        # Configurar interlineado
        if app_instance.formato_config['interlineado'] == 1.0:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        elif app_instance.formato_config['interlineado'] == 1.5:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        if app_instance.formato_config['justificado']:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if app_instance.formato_config['sangria']:
            style.paragraph_format.first_line_indent = Inches(0.5)
        
        style.paragraph_format.space_after = Pt(0)
        
        # Crear estilo de títulos profesional
        try:
            titulo_style = doc.styles.add_style('Titulo Profesional', WD_STYLE_TYPE.PARAGRAPH)
        except:
            titulo_style = doc.styles['Heading 1']
        
        titulo_style.font.name = app_instance.formato_config['fuente_titulo']
        titulo_style.font.size = Pt(app_instance.formato_config['tamaño_titulo'])
        titulo_style.font.bold = True
        titulo_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_style.paragraph_format.space_before = Pt(12)
        titulo_style.paragraph_format.space_after = Pt(12)
        titulo_style.paragraph_format.keep_with_next = True
        titulo_style.paragraph_format.page_break_before = True
        
        return titulo_style
    
    def crear_portada_profesional(self, doc, app_instance):
        """Crea portada profesional con formato de texto mejorado y negrita"""
        # Logo/emblema si existe
        ruta_imagen = self.obtener_ruta_imagen("insignia", app_instance)
        if ruta_imagen and os.path.exists(ruta_imagen):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(ruta_imagen, width=Inches(1.5))
            except Exception as e:
                print(f"Error cargando insignia: {e}")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("LOGO/EMBLEMA DE LA INSTITUCIÓN")
                run.bold = True
                run.font.size = Pt(14)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("LOGO/EMBLEMA DE LA INSTITUCIÓN")
            run.bold = True
            run.font.size = Pt(14)
        
        # Espaciado
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Institución - MEJORADO
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(app_instance.proyecto_data['institucion'].get().upper())
        run.bold = True
        run.font.name = app_instance.formato_config['fuente_titulo']
        run.font.size = Pt(app_instance.formato_config['tamaño_titulo'] + 2)
        
        doc.add_paragraph()
        
        # Título del proyecto - MEJORADO
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'"{app_instance.proyecto_data["titulo"].get()}"')
        run.bold = True
        run.font.name = app_instance.formato_config['fuente_titulo']
        run.font.size = Pt(app_instance.formato_config['tamaño_titulo'] + 4)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Información del proyecto con formato mejorado
        info_fields = [
            ('ciclo', 'Ciclo'),
            ('curso', 'Curso'), 
            ('enfasis', 'Énfasis'),
            ('area', 'Área de Desarrollo'),
            ('categoria', 'Categoría'),
            ('director', 'Director'),
            ('responsable', 'Responsable')
        ]
        
        for field, label in info_fields:
            if field in app_instance.proyecto_data and app_instance.proyecto_data[field].get().strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Crear el texto con etiqueta en negrita y valor normal
                label_run = p.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.name = app_instance.formato_config['fuente_texto']
                label_run.font.size = Pt(app_instance.formato_config['tamaño_texto'])
                
                value_run = p.add_run(app_instance.proyecto_data[field].get())
                value_run.bold = False
                value_run.font.name = app_instance.formato_config['fuente_texto']
                value_run.font.size = Pt(app_instance.formato_config['tamaño_texto'])
        
        # Espaciado adicional
        doc.add_paragraph()
        
        # Estudiantes - MEJORADO
        if app_instance.proyecto_data['estudiantes'].get():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = p.add_run("Estudiantes:")
            title_run.bold = True
            title_run.font.name = app_instance.formato_config['fuente_texto']
            title_run.font.size = Pt(app_instance.formato_config['tamaño_texto'] + 1)
            
            estudiantes = app_instance.proyecto_data['estudiantes'].get().split(',')
            for estudiante in estudiantes:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                student_run = p.add_run(estudiante.strip())
                student_run.font.name = app_instance.formato_config['fuente_texto']
                student_run.font.size = Pt(app_instance.formato_config['tamaño_texto'])
        
        # Tutores - MEJORADO
        if app_instance.proyecto_data['tutores'].get():
            doc.add_paragraph()  # Espaciado
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = p.add_run("Tutores:")
            title_run.bold = True
            title_run.font.name = app_instance.formato_config['fuente_texto']
            title_run.font.size = Pt(app_instance.formato_config['tamaño_texto'] + 1)
            
            tutores = app_instance.proyecto_data['tutores'].get().split(',')
            for tutor in tutores:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tutor_run = p.add_run(tutor.strip())
                tutor_run.font.name = app_instance.formato_config['fuente_texto']
                tutor_run.font.size = Pt(app_instance.formato_config['tamaño_texto'])
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Fecha - MEJORADO
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        year_label = p.add_run("Año: ")
        year_label.bold = True
        year_label.font.name = app_instance.formato_config['fuente_texto']
        year_label.font.size = Pt(app_instance.formato_config['tamaño_texto'])
        
        year_value = p.add_run(str(datetime.now().year))
        year_value.font.name = app_instance.formato_config['fuente_texto']
        year_value.font.size = Pt(app_instance.formato_config['tamaño_texto'])
        
        doc.add_page_break()
    
    def crear_agradecimientos_profesional(self, doc, app_instance):
        """Crea página de agradecimientos con formato profesional"""
        p = doc.add_paragraph()
        p.style = doc.styles['Titulo Profesional']
        run = p.add_run("AGRADECIMIENTOS")
        run.bold = True
        run.font.name = app_instance.formato_config['fuente_titulo']
        run.font.size = Pt(app_instance.formato_config['tamaño_titulo'])
        
        # Configurar como sección de nivel 1
        p.paragraph_format.outline_level = 0
        
        doc.add_paragraph()
        content_p = doc.add_paragraph("(Agregar agradecimientos personalizados aquí)")
        content_p.style = doc.styles['Normal']
        doc.add_page_break()
    
    def crear_indice_profesional(self, doc, app_instance):
        """Crea índice profesional con instrucciones"""
        p = doc.add_paragraph()
        p.style = doc.styles['Titulo Profesional']
        run = p.add_run("ÍNDICE")
        run.bold = True
        run.font.name = app_instance.formato_config['fuente_titulo']
        run.font.size = Pt(app_instance.formato_config['tamaño_titulo'])
        
        # Configurar como sección de nivel 1
        p.paragraph_format.outline_level = 0
        
        doc.add_paragraph()
        
        instrucciones = """INSTRUCCIONES PARA GENERAR ÍNDICE AUTOMÁTICO:

1. En Word, ir a la pestaña "Referencias"
2. Hacer clic en "Tabla de contenido"
3. Seleccionar "Automática" o "Personalizada"
4. El índice se generará automáticamente con todas las secciones

NOTA: Todas las secciones están configuradas con Nivel de esquema 1 para 
facilitar la generación automática del índice."""
        
        content_p = doc.add_paragraph(instrucciones)
        content_p.style = doc.styles['Normal']
        
        doc.add_paragraph()
        
        # Sección para tabla de ilustraciones
        p2 = doc.add_paragraph()
        run2 = p2.add_run("TABLA DE ILUSTRACIONES")
        run2.bold = True
        run2.font.name = app_instance.formato_config['fuente_titulo']
        run2.font.size = Pt(app_instance.formato_config['tamaño_titulo'] - 2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("(Agregar manualmente si hay figuras, tablas o gráficos)")
        doc.add_page_break()
    
    def crear_contenido_dinamico_profesional(self, doc, app_instance):
        """Crea contenido basado en secciones activas con formato profesional"""
        for seccion_id in app_instance.secciones_activas:
            if seccion_id in app_instance.secciones_disponibles:
                seccion = app_instance.secciones_disponibles[seccion_id]
                
                if seccion['capitulo']:
                    # Es un título de capítulo
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(seccion['titulo'].replace('📖 ', '').replace('📚 ', '').replace('🔬 ', '').replace('🛠️ ', '').replace('📊 ', '').replace('💬 ', ''))
                    run.bold = True
                    run.font.name = app_instance.formato_config['fuente_titulo']
                    run.font.size = Pt(app_instance.formato_config['tamaño_titulo'])
                    p.paragraph_format.outline_level = 0
                    p.paragraph_format.page_break_before = True
                    doc.add_paragraph()
                else:
                    # Es contenido
                    if seccion_id in app_instance.content_texts:
                        contenido = app_instance.content_texts[seccion_id].get("1.0", "end").strip()
                        if contenido:
                            titulo_limpio = seccion['titulo']
                            # Remover emojis del título para el documento
                            titulo_limpio = re.sub(r'[^\w\s-]', '', titulo_limpio).strip()
                            self.crear_seccion_profesional(doc, titulo_limpio.upper(), contenido, app_instance)
    
    def crear_seccion_profesional(self, doc, titulo, contenido, app_instance):
        """Crea una sección con formato profesional avanzado"""
        # Título de la sección
        p = doc.add_paragraph()
        p.style = doc.styles['Titulo Profesional']
        run = p.add_run(titulo)
        run.bold = True
        run.font.name = app_instance.formato_config['fuente_titulo']
        run.font.size = Pt(app_instance.formato_config['tamaño_titulo'])
        
        # Configurar como nivel de esquema para índice automático
        p.paragraph_format.outline_level = 0
        p.paragraph_format.keep_with_next = True
        
        # Contenido procesado
        contenido_procesado = self.procesar_citas(contenido.strip(), app_instance)
        if contenido_procesado:
            p = doc.add_paragraph(contenido_procesado)
            p.style = doc.styles['Normal']
            
        doc.add_paragraph()  # Espaciado
    
    def crear_referencias_profesionales(self, doc, app_instance):
        """Crea referencias con formato APA profesional"""
        if not app_instance.referencias:
            return
        
        # Título
        p = doc.add_paragraph()
        p.style = doc.styles['Titulo Profesional']
        run = p.add_run("REFERENCIAS")
        run.bold = True
        run.font.name = app_instance.formato_config['fuente_titulo']
        run.font.size = Pt(app_instance.formato_config['tamaño_titulo'])
        
        # Configurar como sección de nivel 1
        p.paragraph_format.outline_level = 0
        
        doc.add_paragraph()
        
        # Ordenar referencias alfabéticamente
        referencias_ordenadas = sorted(app_instance.referencias, key=lambda x: x['autor'])
        
        for ref in referencias_ordenadas:
            ref_text = f"{ref['autor']} ({ref['año']}). {ref['titulo']}. {ref['fuente']}"
            p = doc.add_paragraph(ref_text)
            # Formato APA: sangría francesa
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            p.style = doc.styles['Normal']
    
    def procesar_citas(self, texto, app_instance):
        """Procesa las citas en el texto"""
        def reemplazar_cita(match):
            cita_completa = match.group(0)
            contenido = cita_completa[6:-1]  # Quita [CITA: y ]
            partes = contenido.split(':')
            
            if len(partes) >= 3:
                tipo, autor, año = partes[0], partes[1], partes[2]
                pagina = partes[3] if len(partes) > 3 else None
                
                if tipo == 'textual':
                    return f" ({autor}, {año}, p. {pagina})" if pagina else f" ({autor}, {año})"
                elif tipo == 'parafraseo':
                    return f" ({autor}, {año})"
                elif tipo == 'larga':
                    return f"\n\n({autor}, {año}, p. {pagina})\n\n" if pagina else f"\n\n({autor}, {año})\n\n"
                elif tipo == 'web':
                    return f" ({autor}, {año})"
                elif tipo == 'multiple':
                    return f" ({autor}, {año})"
            
            return cita_completa
        
        return re.sub(r'\[CITA:[^\]]+\]', reemplazar_cita, texto)
    
    def obtener_ruta_imagen(self, tipo, app_instance):
        """Obtiene la ruta final de la imagen a usar (personalizada o base) - MEJORADA"""
        if tipo == "encabezado":
            # Prioridad: personalizada -> base
            return (getattr(app_instance, 'encabezado_personalizado', None) or 
                   getattr(app_instance, 'ruta_encabezado', None))
        elif tipo == "insignia":
            # Prioridad: personalizada -> base
            return (getattr(app_instance, 'insignia_personalizada', None) or 
                   getattr(app_instance, 'ruta_insignia', None))
        return None
    
    def mostrar_mensaje_exito(self, filename, app_instance):
        """Muestra mensaje de éxito completo"""
        # Mensaje de éxito
        app_instance.validation_text.delete("1.0", "end")
        app_instance.validation_text.insert("1.0", 
            f"🎉 ¡DOCUMENTO PROFESIONAL GENERADO!\n\n"
            f"📄 Archivo: {os.path.basename(filename)}\n"
            f"📍 Ubicación: {filename}\n\n"
            f"✅ CARACTERÍSTICAS APLICADAS:\n"
            f"   • Formato profesional personalizado\n"
            f"   • Niveles de esquema para índice automático\n"
            f"   • Saltos de página entre secciones\n"
            f"   • Control de líneas viudas y huérfanas\n"
            f"   • Conservar títulos con contenido\n"
            f"   • Fuente: {app_instance.formato_config['fuente_texto']} {app_instance.formato_config['tamaño_texto']}pt\n"
            f"   • Títulos: {app_instance.formato_config['fuente_titulo']} {app_instance.formato_config['tamaño_titulo']}pt\n"
            f"   • Imágenes: {'Personalizadas' if getattr(app_instance, 'encabezado_personalizado', None) or getattr(app_instance, 'insignia_personalizada', None) else 'Base'}\n\n"
            f"✅ Estructura profesional ({len(app_instance.secciones_activas)} secciones)\n"
            f"✅ {len(app_instance.referencias)} referencias APA\n"
            f"✅ Citas procesadas automáticamente\n\n"
            f"📋 PARA COMPLETAR EN WORD:\n"
            f"   • Generar índice: Referencias > Tabla de contenido\n"
            f"   • Agregar numeración de páginas si deseas\n"
            f"   • Insertar tablas o figuras según necesidad\n\n"
            f"🚀 ¡Tu proyecto profesional está listo!"
        )
        
        messagebox.showinfo("🎉 ¡Documento Profesional Generado!", 
            f"Documento creado exitosamente:\n{filename}\n\n"
            f"Características aplicadas:\n"
            f"• Formato profesional con niveles de esquema\n"
            f"• Saltos de página automáticos\n"
            f"• Control de líneas profesional\n"
            f"• Imágenes integradas\n"
            f"• Referencias APA automáticas\n\n"
            f"Para generar el índice automático:\n"
            f"Referencias > Tabla de contenido > Automática")